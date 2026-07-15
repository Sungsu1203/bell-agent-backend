"""§paper-writer-1 Step C-2 — end-to-end paper measurement driver.

argparse 7 groups:
  --topic / --sections / --output-dir / --warmup / --measure / --sleep / --timeout / --dry-run

axis 3 metric (B+D 묶음, B-6 design 정합):
  axis 1 (catch 79 재설계): chunk venue OR doi 존재 3등급 실체 판정 — pass_ratio ≥ 0.90,
          3-state verdict(FAIL/WARN/PASS, gate 유지). 구 regex 포화 폐기.
  axis 2: IMRD 4 section — 4 section 존재 + 분량 가이드 ±30% 허용
  axis 3: per-backend ratio (catch 66):
    primary 3: OA ≥ 0.70, SS ≥ 0.40, combined (OA+SS+vertex mean) ≥ 0.50
    보조 metric: vertex_filtered_ratio

표준 (catch 49):
  max_retries=0, provider lock, utf-8 wrapper, stage marker 14단계.

실행:
  .venv_vertex\\Scripts\\python.exe writer_project/scripts/§paper-writer-1/measure_paper.py
  .venv_vertex\\Scripts\\python.exe writer_project/scripts/§paper-writer-1/measure_paper.py --dry-run
"""
from __future__ import annotations

import argparse
import io
import json
import logging
import os
import re
import sys
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any


# ── Step 1: utf-8 wrapper (Windows cp949 회피, catch 49 표준) ──
os.environ["PYTHONIOENCODING"] = "utf-8"
try:
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace", line_buffering=True)
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace", line_buffering=True)
except Exception:
    pass


HERE = Path(__file__).resolve().parent
PROJECT_ROOT = HERE.parents[1]  # writer_project/
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))
# R2-a: scripts/ 를 sys.path 에 얹어 § 없는 중립 패키지(common)를 import 가능하게.
SCRIPTS_ROOT = HERE.parent  # writer_project/scripts
if str(SCRIPTS_ROOT) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_ROOT))

# R2-a: axis3 재설계 — vertex chunk 학술/비학술 가름에 공유 ACADEMIC_DOMAINS 참조.
from urllib.parse import urlparse  # noqa: E402
from common.academic_domains import ACADEMIC_DOMAINS  # noqa: E402


def _chunk_is_academic(chunk: dict) -> bool:
    """chunk 의 domain(우선)·uri(폴백) host 가 ACADEMIC_DOMAINS 에 걸리면 True.

    vertex chunk = {uri, title, domain} (R1 확정) — domain 필드 우선, 없으면 uri host.
    subdomain 허용: host == d 또는 host.endswith('.'+d). www. 프리픽스 정규화.
    """
    host = (chunk.get("domain") or "").strip().lower()
    if not host:
        host = (urlparse(chunk.get("uri") or "").hostname or "").lower()
    host = host.lstrip(".")
    if host.startswith("www."):
        host = host[4:]
    if not host:
        return False
    return any(host == d or host.endswith("." + d) for d in ACADEMIC_DOMAINS)


def _count_backends(chunks: list[dict]) -> dict:
    """chunks 를 _backend 태그로 분류해 카운트 (vertex 는 _chunk_is_academic 로 학술/웹 분해).

    aggregate(_eval_axes)·per-section(_run_one_paper) 공용 — 분류 로직 단일 소스.
    """
    counts = {"openalex": 0, "semantic_scholar": 0,
              "vertex_academic": 0, "vertex_web": 0, "other": 0}
    for c in chunks:
        b = c.get("_backend", "?")
        if b == "openalex":
            counts["openalex"] += 1
        elif b == "semantic_scholar":
            counts["semantic_scholar"] += 1
        elif b == "vertex":
            if _chunk_is_academic(c):
                counts["vertex_academic"] += 1
            else:
                counts["vertex_web"] += 1
        else:
            counts["other"] += 1
    return counts


# ── Step 2: dotenv chain 먼저 (catch 64: override=True 강제) ──
from dotenv import load_dotenv  # noqa: E402
for _env_name in (".env", ".env.vertex", ".env.openalex", ".env.semanticscholar"):
    _p = PROJECT_ROOT / _env_name
    if _p.exists():
        load_dotenv(_p, override=True)


# ── Step 3: driver env override (dotenv 후 강제 — catch 69 후보 lesson) ──
# dotenv override=True 가 driver 사전 설정값을 무효화하는 함정 회피:
# 반드시 dotenv chain 완료 후 driver 가 최종 결정권을 갖도록 강제 set 한다.
# ⚠️ catch 78 정정: SKIP_VERTEX_SEARCH 만은 예외 — config 최초 빌드가 이 함수
#   호출 後(lazy import)라, 토픽 프리셋(topics/<slug>.env override=True)이 최종
#   승자다. 즉 여기 SKIP_VERTEX_SEARCH set 은 CFG 상 no-op(토픽이 결정권).
#   혼동 방지 위해 토픽 값(=1, skip)과 정합되게 "1" 로 맞춰 둔다.
def _force_driver_env_override() -> None:
    """측정 driver 전용 env 강제 override.

    .env 의 LLM_PROVIDER=openai, OPENAI_MAX_RETRIES=1, SKIP_VERTEX_SEARCH=1,
    TOPIC_SLUG=venfobel-vitamin 등이 dotenv override=True 로 driver 사전
    설정값을 덮어쓰는 함정을 차단한다 (catch 69 후보 lesson).
    """
    os.environ["LLM_PROVIDER"] = "vertexai"
    os.environ["VERTEX_MAX_RETRIES"] = "0"
    os.environ["OPENAI_MAX_RETRIES"] = "0"
    os.environ["LLM_MAX_RETRIES"] = "0"
    os.environ["GOOGLE_API_USE_CLIENT_CERTIFICATE"] = "false"
    os.environ["SKIP_VERTEX_SEARCH"] = "1"  # catch 78: no-op(토픽 프리셋 결정) — 토픽 값과 정합
    # ── 맥 SSL: OA·SS urllib HTTPS 검증 (certifi 명시 — macOS 루트 인증서 미연결 회피) ──
    import certifi
    os.environ.setdefault("SSL_CERT_FILE", certifi.where())
    os.environ.setdefault("REQUESTS_CA_BUNDLE", certifi.where())
    os.environ["TOPIC_SLUG"] = "academic-trademark-similarity-consumer"  # "academic-influencer-marketing-consumer-behavior"


_force_driver_env_override()


# ── Step 4: stage marker (timestamp + 번호 + double flush) ──
def _stage(n: int, total: int, desc: str) -> None:
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"[{ts}] [Stage {n}/{total}] {desc}", flush=True)
    sys.stdout.flush()


_stage(1, 14, "dotenv chain + env override 완료")


# ── Step 5: logging ──
logging.basicConfig(level=logging.INFO,
                    format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)


# ── Step 6: lazy imports (env 확정 후 agent 모듈 로드) ──
from agent.web_search import paper_section_fetch  # noqa: E402
from agent.paper_section_writer import (  # noqa: E402
    write_paper_section,
    build_apa_references,
    attach_references_footer,
)

SECTION_WORD_GUIDE = {
    "Introduction": (800, 1200),
    "Theoretical Background": (1000, 1500),
    "Proposed Framework": (1000, 1500),
    "Research Design (Proposed)": (800, 1200),
    "Expected Contributions": (600, 1000),
}

# SECTION_WORD_GUIDE = {
#    "Introduction": (800, 1200),
#    "Methods": (600, 1000),
#    "Results": (1000, 1500),
#    "Discussion": (1500, 2000),
# }

# ── catch 79 종결: axis1 재설계 — regex 포화 → venue OR doi 존재 3등급 실체 판정 ──
# 구 APA_REGEX 는 format_apa7 이 year 를 항상 (YYYY)|(n.d.) 로 뱉어 pass_ratio 1.000
# 포화(변별력 0) = catch 79 함정. 조립 문자열 regex 되파싱(A) 대신 chunk 최상위 필드
# 직독(B)으로 전환. 삭제 아니라 정정 박제(catch 69/70/71 원칙): 후임자 "왜 regex
# 안 쓰나" 오독 방지.  구 regex = r"\(\d{4}\)|\(n\.d\.\)|https?://doi\.org/" (IGNORECASE)


def _blank(v) -> bool:
    """빈 판정: None · '' · 공백 = True (R2.5 지뢰1: SS 는 venue='' 반환).

    str 아닌 truthy(int year 등)는 blank 아님. venue/doi 는 str|None 이라 안전.
    """
    return v is None or (isinstance(v, str) and not v.strip()) or v == ""


def axis1_grade(chunk: dict) -> dict:
    """chunk 최상위 venue/doi 존재로 인용 실체 3등급 판정 (결정론·무료·format-level).

    ⚠️ R2.5 지뢰3: paper_section_fetch 의 OA/SS chunk 는 최상위 필드 형태(metadata
       중첩 아님). 이 함수는 그 최상위 계약을 가정한다.
    사정권 = 존재(가)만. venue 부정합/predatory 품질 판별은 axis1 밖(별 트랙).
    embedding/LLM/네트워크 호출 0.
        완전체 = venue AND doi / 부분체 = venue XOR doi / 결손 = 둘 다 없음(유일 fail).
    doi 필수 아님: 법학 리포지토리 정식인용은 doi 없이도 학술(부분체 pass).
    """
    has_venue = not _blank(chunk.get("venue"))
    has_doi = not _blank(chunk.get("doi"))
    if has_venue and has_doi:
        grade = "complete"
    elif has_venue or has_doi:
        grade = "partial"
    else:
        grade = "missing"
    return {"grade": grade, "pass": grade != "missing",
            "has_venue": has_venue, "has_doi": has_doi}

# ── axis3 재정의 (R2) — 품질 게이트 → 파이프라인 건강/커버리지 기술자 ──
# R1/R1.5: academic_ratio 는 vertex-skip 시 (oa+ss)/(oa+ss)=1.000 포화(변별력 0),
#   유일 런간 변동도 SS 429 노이즈 → 품질 게이트 부적합(catch 79 함정). 게이트 강등.
# 신 역할 = 섹션별 backend 카운트로 degradation 감지(분할 severity):
#   섹션 SS=0 → WARN(429 flaky, 게이트 금지), 섹션 OA=0 → FAIL(완파),
#   other>0 → FAIL(미상 backend 누출 tripwire).
# academic_ratio 는 informational 로만 유지(판정 미사용). 구 임계 상수 폐기.


# ── catch 80: 본문 [[N]] 글로벌 승격 ──
# writer 는 섹션당 로컬 1-based 로 [[N]] 을 매기지만 References footer 는 전 섹션
# concat 글로벌 번호라 섹션 2~5 인용이 오정렬된다. 각 섹션 body 의 [[N]] 에 그 섹션의
# 글로벌 오프셋(= 이전 섹션들 chunk 누적 수)을 더해 footer 와 정합시킨다.
_CITE_MARKER_RE = re.compile(r"\[\[(\d+)\]\]")


def _shift_citation_markers(body: str, offset: int) -> str:
    """body 내 [[N]] 마커를 [[N+offset]] 로 재작성 (자릿수 안전 · offset==0 no-op)."""
    if offset == 0 or not body:
        return body
    return _CITE_MARKER_RE.sub(lambda m: f"[[{int(m.group(1)) + offset}]]", body)


def _build_outline(sections: list[str]) -> str:
    return "\n".join(f"## {i}. {s}" for i, s in enumerate(sections, 1))


def _run_one_paper(topic: str, sections: list[str]) -> dict:
    """1 paper end-to-end run. Returns metrics + body + chunks."""
    outline = _build_outline(sections)
    section_bodies: list[str] = []
    section_chunks_all: list[dict] = []
    per_section: dict[str, dict] = {}
    stage_times: dict[str, float] = {}
    t0 = time.monotonic()

    for i, section in enumerate(sections, 1):
        fetch_stage = 2 + i
        write_stage = 6 + i
        # fetch
        _stage(fetch_stage, 14, f"section {i} ({section}) fetch start")
        tf = time.monotonic()
        chunks = paper_section_fetch(topic, section)
        stage_times[f"fetch_{section}"] = round(time.monotonic() - tf, 2)
        # catch 80: extend 直前 = 이전 섹션들 누적 길이 = 이 섹션의 글로벌 오프셋
        cite_offset = len(section_chunks_all)
        # R2 계측(axis1 재설계): 각 chunk 에 섹션 라벨 태그 (_backend 와 대칭 additive).
        #   build_apa_references·_count_backends·writer 어디도 이 키 미참조 → 로직 영향 0.
        for _c in chunks:
            _c["_section"] = section
        section_chunks_all.extend(chunks)
        # write
        _stage(write_stage, 14, f"section {i} ({section}) LLM 본문 생성")
        tw = time.monotonic()
        target_title = f"{i}. {section}"
        body = write_paper_section(
            topic=topic,
            section_type=section,
            target_title=target_title,
            outline=outline,
            references_chunks=chunks,
            previous_sections="\n\n".join(section_bodies),
            core_thesis=os.environ.get("CORE_THESIS", ""),
        )
        stage_times[f"write_{section}"] = round(time.monotonic() - tw, 2)
        # catch 80: 로컬 [[N]] → 글로벌 [[N+offset]] (footer 정합). offset==0 이면 no-op.
        body = _shift_citation_markers(body, cite_offset)
        section_bodies.append(body)
        per_section[section] = {
            "len_chars": len(body),
            "word_count": len(body.split()),
            "chunks_count": len(chunks),
            "backends": sorted({c.get("_backend", "?") for c in chunks}),
            "backend_counts": _count_backends(chunks),   # P1: 섹션별 backend 분류 (axis3 기술자)
        }

    _stage(11, 14, "References footer build (format_apa7)")
    t_ref = time.monotonic()
    apa_lines = build_apa_references(section_chunks_all)
    paper_body = "\n\n".join(section_bodies)
    paper_full = attach_references_footer(paper_body, apa_lines)
    stage_times["references_footer"] = round(time.monotonic() - t_ref, 2)
    elapsed = time.monotonic() - t0

    return {
        "topic": topic,
        "sections": sections,
        "paper_full": paper_full,
        "section_bodies": section_bodies,
        "per_section": per_section,
        "chunks": section_chunks_all,
        "apa_lines": apa_lines,
        "elapsed_sec": round(elapsed, 2),
        "stage_times": stage_times,
    }


def _eval_axes(result: dict) -> dict:
    """axis 1~3 평가."""
    per_section = result["per_section"]

    # axis 1 (catch 79 재설계): chunk 최상위 venue/doi 존재 3등급 실체 판정.
    #   구 regex(year 포화) 대신 axis1_grade. 입력 = build_apa_references 가 소비하는
    #   그 chunks(최상위 필드). threshold 0.90, 3-state verdict(gate 유지).
    a1_chunks = result["chunks"]
    a1_grades = [axis1_grade(c) for c in a1_chunks]
    a1_n = len(a1_grades)
    a1_dist = {"complete": 0, "partial": 0, "missing": 0}
    for g in a1_grades:
        a1_dist[g["grade"]] += 1
    a1_pass = a1_n - a1_dist["missing"]            # 결손 아닌 것 = pass(venue OR doi)
    a1_ratio = a1_pass / a1_n if a1_n else 0.0
    A1_THRESHOLD = 0.90
    if a1_ratio < A1_THRESHOLD:
        a1_verdict = "FAIL"                        # 결손 과다
    elif a1_dist["missing"] > 0:
        a1_verdict = "WARN"                        # ratio≥th 이나 일부 메타 결손
    else:
        a1_verdict = "PASS"                        # 결손 0
    # 섹션별 결손 (per_section 파생 — _section 태그는 R2 계측이 심음, 한 커밋 의존)
    a1_missing_by_section: dict[str, int] = {}
    for c, g in zip(a1_chunks, a1_grades):
        if g["grade"] == "missing":
            _s = c.get("_section", "?")
            a1_missing_by_section[_s] = a1_missing_by_section.get(_s, 0) + 1
    axis1 = {"pass_ratio": round(a1_ratio, 3), "n": a1_n,
             "threshold": A1_THRESHOLD, "verdict": a1_verdict,
             "grade_dist": a1_dist, "n_missing": a1_dist["missing"],
             "missing_by_section": a1_missing_by_section}

    # axis 2: IMRD section 분량 ±30%
    section_verdicts: dict[str, str] = {}
    for s in result["sections"]:
        lo, hi = SECTION_WORD_GUIDE.get(s, (0, 99999))
        wc = per_section.get(s, {}).get("word_count", 0)
        section_verdicts[s] = "PASS" if int(lo * 0.7) <= wc <= int(hi * 1.3) else "FAIL"
    axis2 = {"per_section": {s: {"word_count": per_section.get(s, {}).get("word_count", 0),
                                  "verdict": v} for s, v in section_verdicts.items()},
             "verdict": "PASS" if all(v == "PASS" for v in section_verdicts.values()) else "FAIL"}

    # axis 3: 파이프라인 건강/커버리지 기술자 (R2 재정의) — 품질 게이트 강등.
    # aggregate 분류는 _count_backends 재사용(per-section 과 단일 소스). 섹션별 신호는
    # per_section[*].backend_counts(P1 캡처)에서 파생. academic_ratio 는 informational 로만.
    chunks = result["chunks"]
    agg = _count_backends(chunks)
    n_oa, n_ss = agg["openalex"], agg["semantic_scholar"]
    n_vx_acad, n_vx_web, n_other = agg["vertex_academic"], agg["vertex_web"], agg["other"]
    total = len(chunks)
    academic_hits = n_oa + n_ss + n_vx_acad          # informational (구 분자, 판정 미사용)
    academic_ratio = academic_hits / total if total else 0.0

    # 섹션별 파생 신호 (P1 이 per_section 에 backend_counts 를 심음)
    per_section_backends = {s: d.get("backend_counts", {}) for s, d in per_section.items()}
    sections_with_zero_ss = [s for s, bc in per_section_backends.items()
                             if bc.get("semantic_scholar", 0) == 0]
    sections_with_zero_oa = [s for s, bc in per_section_backends.items()
                             if bc.get("openalex", 0) == 0]
    sections_with_other = [s for s, bc in per_section_backends.items()
                           if bc.get("other", 0) > 0]

    # 3-state verdict (분할 severity): FAIL(완파/tripwire) > WARN(429 flaky) > PASS.
    if sections_with_zero_oa or sections_with_other:
        a3_verdict = "FAIL"
    elif sections_with_zero_ss:
        a3_verdict = "WARN"
    else:
        a3_verdict = "PASS"

    axis3 = {
        "verdict": a3_verdict,                            # 3-state: PASS / WARN / FAIL
        "sections_with_zero_oa": sections_with_zero_oa,   # FAIL 신호 (OA 결정론 위반 = 완파)
        "sections_with_other": sections_with_other,       # FAIL 신호 (미상 backend tripwire)
        "sections_with_zero_ss": sections_with_zero_ss,   # WARN 신호 (429 flaky, 게이트 아님)
        "per_section_backends": per_section_backends,     # informational (섹션→backend 카운트)
        "backend_counts": {                               # aggregate (vertex 학술/웹 분해)
            "openalex": n_oa, "semantic_scholar": n_ss,
            "vertex_academic": n_vx_acad, "vertex_web": n_vx_web,
            "other": n_other,
        },
        "total": total,
        # informational — 판정 미사용 (R1 에서 품질 게이트 강등)
        "academic_ratio": round(academic_ratio, 3),
        "academic_hits": academic_hits,
    }

    return {"axis1_apa": axis1, "axis2_imrd": axis2, "axis3_pipeline_health": axis3}


def _save_md_docx(result: dict, output_dir: Path, ts: str) -> dict:
    """.md + .docx 양 format 출력."""
    output_dir.mkdir(parents=True, exist_ok=True)
    slug = re.sub(r"[^a-zA-Z0-9_-]+", "_", result["topic"].lower())[:60].strip("_")
    md_path = output_dir / f"paper_{slug}_{ts}.md"
    docx_path = output_dir / f"paper_{slug}_{ts}.docx"

    _stage(12, 14, f".md write → {md_path.name}")
    md_path.write_text(result["paper_full"], encoding="utf-8")

    _stage(13, 14, f".docx write → {docx_path.name}")
    docx_size = 0
    try:
        from docx import Document
        doc = Document()
        for line in result["paper_full"].split("\n"):
            if line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.strip():
                doc.add_paragraph(line)
        doc.save(str(docx_path))
        docx_size = docx_path.stat().st_size
    except Exception as e:
        logger.warning("[docx] failed: %s", e)

    return {
        "md_path": str(md_path), "docx_path": str(docx_path) if docx_size > 0 else "",
        "md_size": md_path.stat().st_size, "docx_size": docx_size,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="§paper-writer-1 measure driver")
    parser.add_argument("--topic", type=str,
                        default=os.environ.get("TOPIC_QUERY"))
    parser.add_argument("--sections", type=str, nargs="+",
                        default=["Introduction", "Theoretical Background",
                                 "Proposed Framework", "Research Design (Proposed)",
                                 "Expected Contributions"])
    parser.add_argument("--output-dir", type=str,
                        default=str(PROJECT_ROOT / "scripts" / "output" / "§paper-writer-1"))
    parser.add_argument("--warmup", type=int, default=0)
    parser.add_argument("--measure", type=int, default=1)
    parser.add_argument("--sleep", type=int, default=30)
    parser.add_argument("--timeout", type=int, default=600)
    parser.add_argument("--dry-run", action="store_true",
                        help="Stage 2 (env dump) 까지만 print 후 exit. API 비용 0.")
    args = parser.parse_args()

    # topic 미정 무음-폴백 박멸: CLI --topic 없음 + env TOPIC_QUERY 없음이면 즉시 에러.
    # 현재 _force_driver_env_override(:131)가 import-time 에 TOPIC_SLUG 를 하드코딩 →
    # config 프리셋(:169)이 TOPIC_QUERY 를 항상 주입하므로 이 가지는 현재 도달 불가.
    # TOPIC_SLUG 동적화/하드코딩 청산 시 살아나는 무음-폴백 안전망이다. 삭제 금지.
    if not args.topic:
        parser.error("--topic 없음 + env TOPIC_QUERY 없음 — topic 미정")

    # Stage 2: config 확정 + env dump
    _stage(2, 14, "config 확정 + env dump")
    for k in ("LLM_PROVIDER", "VERTEX_MAX_RETRIES", "OPENAI_MAX_RETRIES",
              "LLM_MAX_RETRIES", "SKIP_VERTEX_SEARCH", "TOPIC_SLUG",
              "GCP_PROJECT_ID", "GCP_REGION"):
        print(f"  {k}={os.environ.get(k, '(unset)')}", flush=True)
    print(f"  --topic={args.topic!r}", flush=True)
    print(f"  --sections={args.sections}", flush=True)
    print(f"  --timeout={args.timeout}", flush=True)

    if args.dry_run:
        print("\n[--dry-run] Stage 2 완료, exit (API 비용 0).", flush=True)
        return 0

    out_dir = Path(args.output_dir)
    runs: list[dict] = []
    for run_i in range(args.warmup + args.measure):
        phase = "warmup" if run_i < args.warmup else "measure"
        print(f"\n{'='*60}\nrun {run_i + 1}/{args.warmup + args.measure} phase={phase}\n{'='*60}", flush=True)
        r = _run_one_paper(args.topic, args.sections)
        r["phase"] = phase
        r["run_index"] = run_i
        if phase == "measure":
            axes = _eval_axes(r)
            r.update(axes)
            ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
            files = _save_md_docx(r, out_dir, ts)
            r["files"] = files
            _stage(14, 14, "axis 측정 완료 + 결과 박제")
        runs.append(r)
        if run_i < args.warmup + args.measure - 1:
            time.sleep(args.sleep)

    # JSON 박제 (body/chunks 제외)
    measurements = [
        {k: v for k, v in r.items() if k not in ("paper_full", "section_bodies", "chunks")}
        for r in runs
    ]
    summary = {
        "topic": args.topic, "sections": args.sections,
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "warmup": args.warmup, "measure": args.measure,
        "runs": measurements,
    }
    json_path = out_dir / "c_paper_measurement.json"
    out_dir.mkdir(parents=True, exist_ok=True)
    json_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\n[saved] {json_path}", flush=True)

    # ── R2 계측(axis1 재설계): chunks 원본 별도 덤프 (R3 detector 오프라인 튜닝 재료) ──
    # c_paper_measurement.json 오염 방지 위해 별도 파일. 실체 4필드(authors/year/venue/doi)
    # + title + backend/section 태그만 추림. axis1 채점 로직 무관 = 순수 additive 계측.
    chunks_dump: list[dict] = []
    for r in runs:
        if r.get("phase") != "measure":
            continue
        for c in r.get("chunks", []) or []:
            chunks_dump.append({
                "section": c.get("_section"),
                "backend": c.get("_backend"),
                "title": c.get("title"),
                "authors": c.get("authors"),
                "year": c.get("year"),
                "venue": c.get("venue"),
                "doi": c.get("doi"),
            })
    dump_path = out_dir / "chunks_raw_dump.json"
    dump_path.write_text(json.dumps(
        {"generated_at_utc": summary["generated_at_utc"],
         "topic": args.topic, "n_chunks": len(chunks_dump),
         "chunks": chunks_dump}, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[saved] {dump_path} ({len(chunks_dump)} chunks)", flush=True)

    print(json.dumps(measurements[-1] if measurements else {}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
