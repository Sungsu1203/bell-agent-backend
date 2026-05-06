# app.py (final)

from __future__ import annotations

import io
import re
from rag_expression import (
    extract_new_topic_title,
    extract_section_index,
    extract_write_title,
    is_outline_creation,
    is_outline_display,
)

from core.models import Task

from dataclasses import dataclass
from pathlib import Path
import re
from typing import Any, Dict, Optional, Tuple
from collections.abc import MutableMapping

# ✅ repo rule 기반 outline 경로 생성 함수
from core.paths import outline_path, _default_outline_name

from fastapi import HTTPException
from fastapi.responses import StreamingResponse

# ── Fault handler (deadlock/멈춤 추적) ──────────────────────────────────────────
import faulthandler, sys, time
faulthandler.enable()  # 항상 가능한 초반에 활성화
# 타임아웃(초) 환경변수로 제어: FAULT_DUMP_AFTER_SEC (기본 180초)
import os as _os_for_fault
TIMEOUT_DUMP_SEC = int((_os_for_fault.getenv("FAULT_DUMP_AFTER_SEC") or "180").strip() or "180")
# 콘솔(stderr)로 타임아웃 스택 덤프 예약은 유지(파일 예약은 로깅 설정 후에 별도 파일로 수행)
faulthandler.dump_traceback_later(TIMEOUT_DUMP_SEC, file=sys.stderr)

_trace_fh = None  # 파일 트레이스 핸들(로깅 설정 후 열어 예약)

def _cancel_fault_timers_and_close() -> None:
    """예약된 스택 덤프 타이머를 취소하고 파일 핸들을 정리합니다."""
    try:
        faulthandler.cancel_dump_traceback_later()
    except Exception:
        pass
    try:
        global _trace_fh
        if _trace_fh and not _trace_fh.closed:
            _trace_fh.close()
            _trace_fh = None
    except Exception:
        pass

def _fault_log_path(main_log_file: str) -> str:
    """
    faulthandler 전용 파일 경로를 돌려줍니다.
    - 환경변수 FAULT_LOG_FILE 로 강제 지정 가능
    - 기본: main_log_file과 같은 폴더의 'run_stack.log'
    """
    override = (_os_for_fault.getenv("FAULT_LOG_FILE") or "").strip()
    if override:
        return override
    import os as _os
    d = _os.path.dirname(main_log_file) or "."
    return _os.path.join(d, "run_stack.log")


# ── 환경 로드 ───────────────────────────────────────────────────────────────────
import os

# ── Early debug buffer & mode ──────────────────────────────────────────────────
def _env_truthy(name: str, default: bool = False) -> bool:
    v = (os.getenv(name) or "").strip().lower()
    if not v:
        return default
    return v in ("1", "true", "yes", "on")

DEBUG_MODE = _env_truthy("DEBUG", False) or (os.getenv("LOG_LEVEL", "INFO").strip().upper() == "DEBUG")
_EARLY_DEBUG: list[str] = []

def early_debug(msg: str) -> None:
    """로깅 설정 전 수집. DEBUG 모드에서만 나중에 방출."""
    if DEBUG_MODE:
        _EARLY_DEBUG.append(msg)

from dotenv import load_dotenv, find_dotenv
dotenv_path = find_dotenv(filename=".env", usecwd=True)
# 1) .env 로드만 수행(주입/강제 덮어쓰기 제거) → 단일 진입점은 CFG에서 처리
load_dotenv(dotenv_path=dotenv_path, override=False)

# ── ENV DIAGNOSTIC (safe: only length) ─────────────────────────
def _mask_len(v: str | None) -> int:
    return len(v.strip()) if v else 0

# ── 표준 라이브러리 ────────────────────────────────────────────────────────────
import sys
import argparse
from logging import StreamHandler, Formatter
import logging
from typing import Optional, TextIO, Any, Dict, Tuple, cast
from logging.handlers import RotatingFileHandler
import json

import certifi

from pydantic import BaseModel

os.environ["SSL_CERT_FILE"] = certifi.where()
os.environ["REQUESTS_CA_BUNDLE"] = certifi.where()

# ── Web server (FastAPI) ──────────────────────────────────────────────────────
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
import asyncio
import threading

from collections import deque
from fastapi.responses import FileResponse
from pathlib import Path
import threading

# ── 프로젝트 의존 ───────────────────────────────────────────────────────────────
from content.api import find_section_path  # re-export 사용(섹션 파일 탐색)
from langchain_core.messages import SystemMessage, HumanMessage, AIMessage
import core.config as config

from core.paths import now_str as _now_str, current_path
from core.state_types import State

from core.state_io import save_state
from core.events import emit_event, get_events_since, clear_events
from utils.sanitize import coerce_int
from report_builder import build_final_report

_ALLOWED_FLAG_KEYS = {
    "pending_write_title","requested_write_title","suppress_vector_qa",
    "sections_done","sections_total","sections_seen",
    "chapters_done","chapters_total","chapters_seen",
    "dash_last_ts","dash_count","skip_web_search","debug",
    "topic_title","iteration_count",
}

# Flags TypedDict의 필수 키를 만족시키기 위한 기본값 테이블
# (필수/선택 구분과 무관하게 넉넉히 채워 안전 캐스팅)
_DEFAULT_FLAGS_SHAPE: Dict[str, Any] = {
    "pending_write_title": False,
    "requested_write_title": "",
    "suppress_vector_qa": False,
    "sections_done": 0,
    "sections_total": 0,
    "sections_seen": "",
    "chapters_done": 0,
    "chapters_total": 0,
    "chapters_seen": "",
    "dash_last_ts": "",
    "dash_count": 0,
    "skip_web_search": False,
    "debug": False,
    "topic_title": "",
    "iteration_count": 0,
}


def update_flags(state: State, **updates: Any) -> None:
    """Flags에 허용된 키만 반영하고, 나머지는 state 루트에 기록."""
    f_raw = state.get("flags") or {}
    # dict로 복사(원본이 TypedDict 이어도 가변 사본으로 작업)
    f: Dict[str, Any] = dict(f_raw)

    # 1) 업데이트 반영: 허용 키만 flags로 기록 (루트에는 동적 키를 쓰지 않음)
    for k, v in updates.items():
        if k in _ALLOWED_FLAG_KEYS:
            f[k] = v
        # else: 무시 (루트에 동적 키 쓰면 mypy 에러)

    # 2) Flags TypedDict의 shape 보장(필수 키 기본값 채움)
    for k, v in _DEFAULT_FLAGS_SHAPE.items():
        f.setdefault(k, v)

    # 3) 최종 대입(형 안전 캐스팅)
    state["flags"] = cast(Dict[str, Any], f)

if bool(getattr(config.CFG, "HUMAN_LOGS_VERBOSE", False)):
    print(config.CFG)  # CFG를 단일 소스로 노출(디버그 시)

# ── 로깅 설정 ──────────────────────────────────────────────────────────────────
def _int_env(name: str, default: int) -> int:
    try:
        v = (os.getenv(name) or "").strip()
        return int(v) if v != "" else default
    except Exception:
        return default

def _bool_env(name: str, default: bool = False) -> bool:
    # CFG 우선 → 없으면 ENV 폴백(점진적 호환)
    try:
        if hasattr(config.CFG, name):
            return bool(getattr(config.CFG, name))
    except Exception:
        pass
    v = (os.getenv(name) or "").strip().lower()
    if not v: return default
    return v in ("1","true","yes","on")

def _truthy_cfg(name: str, fallback_env: str | None = None, default: bool = False) -> bool:
    if hasattr(config.CFG, name):
        return bool(getattr(config.CFG, name))
    return (os.getenv(fallback_env or name, "1" if default else "0").strip().lower() in ("1","true","yes","on"))

def human_print(msg: str):
    if bool(getattr(config.CFG, "HUMAN_LOGS", False)):
        print(msg, flush=True)

def setup_logging(stream: Optional[TextIO] = None) -> None:
    root = logging.getLogger()

    for h in list(root.handlers):
        root.removeHandler(h)
        try:
            h.close()
        except Exception:
            pass

    # LOG_LEVEL 등은 CFG를 단일 진입점으로 사용
    base_level_name = (str(getattr(config.CFG, "LOG_LEVEL", "INFO")) or "INFO").strip().upper()
    console_level_name = (str(getattr(config.CFG, "LOG_LEVEL_CONSOLE", base_level_name)) or base_level_name).strip().upper()
    file_level_name    = (str(getattr(config.CFG, "LOG_LEVEL_FILE",    base_level_name)) or base_level_name).strip().upper()

    # root logger는 핸들러 중 "더 낮은 레벨"로 설정 (DEBUG 파일 + INFO 콘솔을 가능하게)
    base_level = getattr(logging, base_level_name, logging.INFO)
    console_level = getattr(logging, console_level_name, base_level)
    file_level = getattr(logging, file_level_name, base_level)
    root_level = min(console_level, file_level)
    root.setLevel(root_level)
    root.propagate = False

    logging.getLogger("httpx").setLevel(logging.WARNING)
    logging.getLogger("httpcore").setLevel(logging.WARNING)
    logging.getLogger("chromadb").setLevel(logging.WARNING)
    logging.getLogger("urllib3").setLevel(logging.WARNING)
    logging.getLogger("google").setLevel(logging.WARNING)
    logging.getLogger("unstructured").setLevel(logging.WARNING)
    logging.getLogger("absl").setLevel(logging.WARNING)
    logging.getLogger("grpc").setLevel(logging.WARNING)

    use_json = bool(getattr(config.CFG, "LOG_JSON", False))
    fmt = getattr(config.CFG, "LOG_FMT", "[%(levelname)s] %(name)s: %(message)s")
    datefmt = getattr(config.CFG, "LOG_DATEFMT", "%Y-%m-%dT%H:%M:%S")

    if use_json:
        class _JsonFormatter(logging.Formatter):
            def format(self, record: logging.LogRecord) -> str:
                obj = {
                    "time": self.formatTime(record, datefmt),
                    "level": getattr(record, "levelname", ""),
                    "name": record.name,
                    "module": record.module,
                    "lineno": record.lineno,
                    "message": record.getMessage(),
                }
                if record.exc_info:
                    obj["exc_info"] = self.formatException(record.exc_info)
                return json.dumps(obj, ensure_ascii=False)
        formatter: logging.Formatter = _JsonFormatter()
    else:
        formatter = logging.Formatter(fmt=fmt, datefmt=datefmt)

    import re as _re
    class HumanOnlyFilter(logging.Filter):
        _pat_basic = _re.compile(
            r'^\s*User\s*:'
            r'|^\s*\[REPORT\]'
            r'|^\s*Application started'
            r'|^\s*Goodbye!'
            r'|^\s*MESSAGE COUNT'
        , _re.I)

        def filter(self, record: logging.LogRecord) -> bool:
            msg = str(record.getMessage() or "")

            # 콘솔에서 부팅/환경 배너 숨김 (파일 로그에는 그대로 남김)
            # - app.py에서 logger.info로 찍는 시작 배너류
            if msg.startswith(("ENV TOPIC_TITLE=", "ENV CHECK:", "[GATEKEEP]", "Application started")):
                return False

            if (os.getenv("HUMAN_LOGS_STRICT","0").lower() in ("1","true","yes","on")):
                return bool(self._pat_basic.search(msg))
            if record.name.startswith(("agent.", "tools.", "utils.")):
                return False
            if self._pat_basic.search(msg):
                return True
            if record.levelno >= logging.INFO and record.name in {"__main__", "report_builder"}:
                return True
            return False

    show_human_only = bool(getattr(config.CFG, "HUMAN_LOGS", False))

    ch = logging.StreamHandler(stream or sys.stdout)
    ch.setLevel(console_level)
    if show_human_only:
        ch.addFilter(HumanOnlyFilter())
        ch.setFormatter(logging.Formatter("%(message)s"))
    else:
        ch.setFormatter(formatter)
    root.addHandler(ch)

    # (선택) DEBUG 모드가 아닐 때 [CRITICAL DEBUG] 태그 숨김
    if not DEBUG_MODE:
        class _DropCriticalDebug(logging.Filter):
            def filter(self, record: logging.LogRecord) -> bool:
                try:
                    return "[CRITICAL DEBUG]" not in record.getMessage()
                except Exception:
                    return True
        root.addFilter(_DropCriticalDebug())

    # 파일 로거 설정도 CFG 우선
    log_file = (getattr(config.CFG, "LOG_FILE", "") or "").strip() or os.path.join(".", "logs", "run_full.log")
    max_bytes = int(getattr(config.CFG, "LOG_MAX_BYTES", 1048576) or 1048576)
    backups = int(getattr(config.CFG, "LOG_BACKUP_COUNT", 3) or 3)

    log_dir = os.path.dirname(log_file)
    if log_dir:
        os.makedirs(log_dir, exist_ok=True)

    fh = RotatingFileHandler(log_file, maxBytes=max_bytes, backupCount=backups, encoding="utf-8", delay=True)
    fh.setLevel(file_level)
    fh.setFormatter(formatter)
    root.addHandler(fh)

    # ▼ faulthandler: 파일로도 덤프 예약(로깅 경로 확정 후)
    global _trace_fh
    try:
        # 중요: 로거 파일과 '다른 파일'을 사용해야 Windows에서 롤오버 충돌(WinError 32)을 피할 수 있습니다.
        fault_file = _fault_log_path(log_file)
        os.makedirs(os.path.dirname(fault_file) or ".", exist_ok=True)
        _trace_fh = open(fault_file, "a", encoding="utf-8", buffering=1)
        # 콘솔 예약은 이미 설정됨 → 파일로 추가 예약(마지막 예약이 유효)
        faulthandler.dump_traceback_later(TIMEOUT_DUMP_SEC, file=_trace_fh)
        # 초기 안내(디버그)
        if bool(getattr(config.CFG, "HUMAN_LOGS_VERBOSE", False)):
            logging.getLogger("app.init").debug("[FAULT] file dump scheduled → %s", fault_file)
    except Exception:
        _trace_fh = None
    # ▲

logger = logging.getLogger(__name__)

# ── Web server globals ────────────────────────────────────────────────────────
web_app = FastAPI(title="RAG Writer API", version="0.1")

# Next.js(기본 3000)에서 호출 가능하도록 CORS 허용
web_app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://127.0.0.1:3000",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"],   # ← 이 줄 추가
)

# 한 번에 한 요청만 그래프를 실행(상태 공유 + 안전)
_RUN_LOCK = asyncio.Lock()
_CANCEL_EVENT = threading.Event()

# 웹 모드에서 사용할 런타임 상태(프로세스 메모리)
_WEB_STATE: Optional[State] = None
_WEB_ARGS: Any = None  # argparse args

# ── Web logs buffer (minimal polling) ─────────────────────────────────────────
_LOG_MAX = int(os.getenv("WEB_LOG_MAX", "2000"))  # 메모리 로그 최대 줄 수
_LOG_BUF = deque(maxlen=_LOG_MAX)                 # (seq:int, line:str)
_LOG_SEQ = 0
_LOG_LOCK = threading.Lock()

class _WebLogHandler(logging.Handler):
    """서버 로그를 메모리 버퍼에 쌓아 /api/logs로 제공."""
    def emit(self, record: logging.LogRecord) -> None:
        global _LOG_SEQ
        try:
            msg = self.format(record)
        except Exception:
            msg = str(getattr(record, "msg", "") or "")
        with _LOG_LOCK:
            _LOG_SEQ += 1
            _LOG_BUF.append((_LOG_SEQ, msg))

_web_log_handler: logging.Handler | None = None

def attach_web_log_handler() -> None:
    """setup_logging() 이후 1회 호출 권장."""
    global _web_log_handler
    if _web_log_handler is not None:
        return
    h = _WebLogHandler()
    # 파일 로그/콘솔과 별개로 “메시지 자체”만 저장하면 되므로 단순 포맷
    h.setFormatter(logging.Formatter("%(asctime)s %(levelname)s [%(name)s] %(message)s"))
    logging.getLogger().addHandler(h)
    _web_log_handler = h

# ── Path helpers (fit to your folder structure) ──────────────────────────────
def _topic_slug_from_state(st: State) -> str:
    return str(st.get("topic_slug") or "default").strip() or "default"

def _outline_fname_from_state(st: State) -> str:
    # report 모드면 outline_report.md가 기본
    return str(st.get("outline_fname") or ("outline_report.md" if config.DOC_MODE == "report" else "outline.md"))

def _outline_paths(st: State) -> tuple[str, str]:
    """
    선생님 구조:
      outlines/<slug>/<outline_fname> (우선)
      outlines/default/<outline_fname> (폴백)
    """
    slug = _topic_slug_from_state(st)
    fname = _outline_fname_from_state(st)
    p1 = os.path.join(str(current_path), "outlines", slug, fname)
    p2 = os.path.join(str(current_path), "outlines", "default", fname)
    return p1, p2

def _doc_mode_from_cfg() -> str:
    # CFG가 우선. 혹시 None이면 report로.
    return str(getattr(config.CFG, "DOC_MODE", "report") or "report").strip().lower()

def _artifact_root_for_mode(mode: str) -> str:
    # 중요한 산출물 루트 폴더
    return "chapters" if mode == "book" else "sections"

def _artifact_dir(mode: str, slug: str) -> str:
    return os.path.join(str(current_path), _artifact_root_for_mode(mode), slug)

def _safe_under_artifacts(path: str) -> bool:
    """
    sections/ 또는 chapters/ 또는 reports/ 아래만 다운로드 허용 (경로 탈출 방지)
    """
    try:
        base_sections = Path(os.path.join(str(current_path), "sections")).resolve()
        base_chapters  = Path(os.path.join(str(current_path), "chapters")).resolve()
        base_reports   = Path(os.path.join(str(current_path), "reports")).resolve()
        target = Path(path).resolve()
        s = str(target)
        return (
            s.startswith(str(base_sections)) or
            s.startswith(str(base_chapters)) or
            s.startswith(str(base_reports))
        )
    except Exception:
        return False



# ── 도움말 ──────────────────────────────────────────────────────────────────────
def print_help():
    logger.info(
        "\n[도움말]\n"
        "- 일반 입력: 에이전트에게 지시/질문을 보냅니다.\n"
        "- 여러 줄 입력: 첫 줄에 ``` 또는 \"\"\" 를 입력해 시작/종료합니다.\n"
        "- 줄바꿈 이어쓰기: 줄 끝에 \\ 입력\n"
        "- 종료: exit | quit | q\n"
        "- 예시: '목차 생성', '최신 자료로 RAG 업데이트', 'write: 서론'\n"
    )

# ── 초기 상태 ──────────────────────────────────────────────────────────────────
def initial_state(iteration_count: int, agent_role: str | None = None) -> State:
    default_outline = "outline_report.md" if config.DOC_MODE == "report" else "outline.md"
    # mypy: 초기 dict는 동적 키가 섞이므로 Dict로 생성 후 마지막에 cast(State)
    # 초기 구성은 동적 키가 섞이므로 Dict로 만들고 마지막에 State로 캐스팅
    base: Dict[str, Any] = {
        "messages": [SystemMessage(content=(
            f"너희 AI들은 사용자의 요구에 맞는 {('보고서' if config.DOC_MODE=='report' else '책')}을(를) 쓰는 작가팀이다. "
            f"사용자가 사용하는 언어로 대화하라. 현재시각은 {_now_str()}이다. "
            f"항상 한국어로 작성하라. 사용자에게서 한국어/영어가 섞여와도 산출물은 전부 한국어로 통일하라."
        ))],
        "task_history": [],
        "references": {"queries": [], "docs": []},
        # 일부 라우터/도구는 경량 미러 키 'refs'도 참조하므로 초깃값에서 함께 준비
        "refs": {"queries": [], "docs": []},
        # Optional/Any 섞임 방지: 안전한 문자열 변환 후 처리
        # CFG.AGENT_ROLE 우선(하위호환: 호출 인자/BLOCKAGI_*는 상위에서 폴백)
        "agent_role": str(agent_role or getattr(config.CFG, "AGENT_ROLE", "") or "").strip().lower(),
        "iteration_count": int(iteration_count),
        "research_objectives": [],
        "research_round": 0,
        "research_loop_active": False,
        "findings_md": [],
        "llm_logs": [],
        "new_url_count": None,
        "topic_slug": getattr(config.CFG, "TOPIC_SLUG", "") or "default",
        "outline_fname": default_outline,
        "outline_shown": False,
        "facts_ctx": "",
        "research_plan": {"round": 0, "objective": "", "queries": [], "timestamp": _now_str()},
        "flags": {},
    }
    return cast(State, base)

# ── 사용자 입력 ────────────────────────────────────────────────────────────────
def read_user_input() -> str:
    try:
        # first = input("\nUser\t: ")
        first = input("\n> ")
    except (EOFError, KeyboardInterrupt):
        raise
    s = first.strip()

    if s.lower() in ("help", "?"):
        print_help()
        return ""

    if s.startswith('```') or s == '"""':
        fence = '```' if s.startswith('```') else '"""'
        lines: list[str] = []
        while True:
            try:
                line = input()
            except (EOFError, KeyboardInterrupt):
                return "\n".join(lines).strip()
            if line.strip() == fence:
                break
            lines.append(line)
        return "\n".join(lines).strip()

    buf = first
    while buf.endswith("\\"):
        try:
            buf = buf[:-1] + "\n" + input()
        except (EOFError, KeyboardInterrupt):
            break
    return buf.strip()


# ── 지연 로딩: 그래프 빌더 ─────────────────────────────────────────────────────
def _load_graph():
    from graph import build_graph  # InvokableGraph 타입이 있더라도 여기선 불필요
    return build_graph()

# ── 그래프 캐시 ────────────────────────────────────────────────────────────────
# 모듈 전역 캐시(타입 힌트 포함)
_graph_obj: Optional[Any] = None
_graph_sig: Optional[Tuple[Any, ...]] = None

def _graph_signature() -> Tuple[Any, ...]:
    """
    그래프 재빌드 여부를 결정할 '구성 시그니처'.
    CFG가 바뀌면 여기 구성 요소를 더 넣으세요.
    """
    try:
        cfg = config.CFG
        return (
            getattr(config, "DOC_MODE", "report"),
            bool(getattr(cfg, "ENABLE_COMMUNICATOR", True)),
            bool(getattr(cfg, "ENABLE_CONTENT_STRATEGIST", True)),
            bool(getattr(cfg, "ENABLE_VECTOR_SEARCH", True)),
            bool(getattr(cfg, "ENABLE_WEB_SEARCH", True)),
            bool(getattr(cfg, "ENABLE_CHAPTER_WRITER", True)),
            bool(getattr(cfg, "ENABLE_SECTION_WRITER", True)),
            bool(getattr(cfg, "ENABLE_RESEARCH_PLANNER", True)),
            bool(getattr(cfg, "ENABLE_RESEARCH_SYNTHESIZER", True)),
        )
    except Exception:
        # 문제가 생기면 '항상 동일'한 튜플을 반환해 과도한 리빌드를 막음
        return ("fallback",)

def _get_graph_cached():
    global _graph_obj, _graph_sig
    sig = _graph_signature()
    if _graph_obj is None or _graph_sig != sig:
        g = _load_graph()  # graph.build_graph()를 내부에서 호출
        if g is None or not hasattr(g, "invoke"):
            raise RuntimeError("Graph build failed: compiled graph has no 'invoke'.")
        _graph_obj, _graph_sig = g, sig
        logger.debug("[GRAPH] (re)built; signature=%s", sig)
    return _graph_obj

def _last_ai_text(state: State) -> str:
    """state.messages에서 마지막 AIMessage의 content만 뽑아 텍스트로 반환."""
    try:
        from langchain_core.messages import AIMessage
        msgs = list(state.get("messages", []))
        last_ai = next((m for m in reversed(msgs) if isinstance(m, AIMessage)), None)
        return str(getattr(last_ai, "content", "") or "").strip() if last_ai else ""
    except Exception:
        return ""

def parse_command_intent(text: str) -> dict:
    s = (text or "").strip()
    s_l = s.lower()

    # 1) force_query: inline
    m = re.match(r"^\s*force_query\s*:\s*(.+?)\s*$", s, flags=re.IGNORECASE)
    if m:
        q = (m.group(1) or "").strip().strip('"').strip("'")
        return {"type": "force_queries", "payload": {"queries": [q] if q else []}, "raw": s[:160]}

    # 2) build report는 run_once에서 이미 특수처리 하므로 표시만 해도 됨(선택)
    if s_l in ("build: report", "report build", "보고서 빌드", "최종 보고서 생성"):
        return {"type": "build_report", "payload": {}, "raw": s[:160]}

    # 3) 새 주제
    try:
        t = extract_new_topic_title(s)
    except Exception:
        t = None
    if t:
        return {"type": "new_topic", "payload": {"title": t}, "raw": s[:160]}

    # 4) 목차 show/create
    try:
        if is_outline_display(s):
            return {"type": "show_outline", "payload": {}, "raw": s[:160]}

        if is_outline_creation(s):
            return {"type": "create_outline", "payload": {}, "raw": s[:160]}
    except Exception:
        pass

    # 5) RAG update (supervisor와 동일한 정규식)
    _rag_re = r"(최신|업데이트|update|latest).*?(자료|리소스|레퍼런스|참고|sources|material).*?(rag|벡터|vector|임베딩|embedding|index|색인|chroma)"
    if re.search(_rag_re, s, flags=re.IGNORECASE):
        return {"type": "rag_update", "payload": {"mode": "auto"}, "raw": s[:160]}

    # 6) write:
    try:
        w = extract_write_title(s)
    except Exception:
        w = None
    if not w:
        try:
            idx = extract_section_index(s)
        except Exception:
            idx = None
        if idx:
            return {"type": "write_index", "payload": {"index": idx}, "raw": s[:160]}
    if w:
        return {"type": "write", "payload": {"title": w}, "raw": s[:160]}

    return {"type": "none", "payload": {}, "raw": s[:160]}

def _has_pending_task(state: State, agent: str, prefix: str | None = None) -> bool:
    for t in (state.get("task_history") or []):
        try:
            if getattr(t, "done", False):
                continue
            if getattr(t, "agent", "") != agent:
                continue
            if prefix:
                desc = str(getattr(t, "description", "") or "")
                if not desc.startswith(prefix):
                    continue
            return True
        except Exception:
            continue
    return False



@dataclass(frozen=True)
class FastResponse:
    ok: bool
    mode: str              # always "fast"
    kind: str              # "outline"|"status"|...
    message: str
    last_saved_path: Optional[str] = None
    state_patch: Optional[Dict[str, Any]] = None
    artifacts: Optional[Dict[str, Any]] = None


class FastCommandHandler:
    """
    Deterministic only. NEVER calls graph/supervisor/LLM/embedding/chroma.
    """

    # Natural-language aliases allowed in fast (keep small & conservative)
    _OUTLINE_ALIASES = {"목차", "목차 보여줘", "목차보여줘", "아웃라인"}

    # Slash commands
    _RE_SLASH = re.compile(r"^\s*/(?P<cmd>[a-zA-Z0-9_-]+)(?:\s+(?P<arg>.*?))?\s*$")
    _RE_SET = re.compile(r"^\s*(?P<key>[A-Z0-9_]+)\s*=\s*(?P<val>.+?)\s*$")

    @staticmethod
    def _topic_slug(state: MutableMapping[str, Any]) -> Optional[str]:
        slug = state.get("topic_slug") or state.get("topic") or state.get("topicTitle")
        if not isinstance(slug, str):
            return None
        slug = slug.strip()
        return slug or None

    @staticmethod
    def _doc_mode(state: MutableMapping[str, Any]) -> str:
        """
        Deterministic doc mode resolver.
        우선순위:
          1) state["flags"]["DOC_MODE"]
          2) state["doc_mode"]
          3) default: "report"
        """
        flags = state.get("flags")
        if isinstance(flags, dict):
            v = flags.get("DOC_MODE")
            if isinstance(v, str):
                vv = v.strip().lower()
                if vv in ("report", "book"):
                    return vv

        v2 = state.get("doc_mode")
        if isinstance(v2, str):
            vv2 = v2.strip().lower()
            if vv2 in ("report", "book"):
                return vv2

        return "report"

    @staticmethod
    def DEFAULT_OUTLINE_PATH(state: MutableMapping[str, Any], topic_slug: str) -> Path:
        """
        ✅ 레포 규칙 기반 단일 경로 확정:
        - filename: core.paths._default_outline_name(mode)
        - path join: core.paths.outline_path(filename, topic_slug=..., mode=...)
        """
        mode = FastCommandHandler._doc_mode(state)
        if mode not in ("report", "book"):
            mode = "report"
        fname = _default_outline_name(mode)
        # outline_path는 보통 str을 반환하므로 Path로 감싼다
        return Path(outline_path(fname, topic_slug=topic_slug, mode=mode))

    @staticmethod
    def _truncate(text: str, max_chars: int = 8000) -> str:
        text = text.strip()
        if len(text) <= max_chars:
            return text
        return text[:max_chars].rstrip() + "\n\n...(생략: 너무 길어서 일부만 표시)"

    @staticmethod
    def _apply_patch(state: MutableMapping[str, Any], patch: Dict[str, Any]) -> None:
        """
        Shallow apply for top-level keys and nested dict under "flags".
        Deterministic, minimal.
        """
        for k, v in patch.items():
            if k == "flags" and isinstance(v, dict):
                flags = state.get("flags")
                if not isinstance(flags, dict):
                    flags = {}
                    state["flags"] = flags
                flags.update(v)
            else:
                state[k] = v

    def handle(self, text: str, state: MutableMapping[str, Any]) -> Tuple[bool, Optional[FastResponse]]:
        raw = (text or "").strip()
        if not raw:
            return False, None

        # 1) Natural-language outline aliases -> fast outline
        if raw in self._OUTLINE_ALIASES:
            return True, self._handle_outline(state)

        # 2) Slash commands
        m = self._RE_SLASH.match(raw)
        if not m:
            return False, None

        cmd = (m.group("cmd") or "").lower()
        arg = (m.group("arg") or "").strip()

        if cmd in ("outline", "ol"):
            return True, self._handle_outline(state)

        if cmd in ("status", "st"):
            return True, self._handle_status(state)

        if cmd in ("help", "h"):
            return True, self._handle_help()

        if cmd in ("new-topic", "topic", "nt"):
            return True, self._handle_new_topic(arg, state)

        if cmd in ("set",):
            return True, self._handle_set(arg, state)

        if cmd in ("write", "w"):
            return True, self._handle_write_queue(arg, state)

        # Unknown slash command -> fast help (still deterministic)
        return True, FastResponse(
            ok=False,
            mode="fast",
            kind="error",
            message=f"알 수 없는 명령: /{cmd}\n\n/help 로 사용 가능한 명령을 확인하세요.",
            state_patch={"execution_mode": "fast"},
        )

    def _handle_outline(self, state: MutableMapping[str, Any]) -> FastResponse:
        slug = self._topic_slug(state)
        if not slug:
            patch = {"execution_mode": "fast"}
            self._apply_patch(state, patch)
            return FastResponse(
                ok=False,
                mode="fast",
                kind="outline",
                message="topic_slug가 설정되지 않았습니다. 먼저 /new-topic <slug> 를 실행하세요.",
                state_patch=patch,
            )

        # ✅ core/paths.py 규칙으로 '단일 경로' 확정
        outline_path_p = self.DEFAULT_OUTLINE_PATH(state, slug)

        if not outline_path_p.exists():
            patch = {"execution_mode": "fast", "outline_path": str(outline_path_p)}
            self._apply_patch(state, patch)
            return FastResponse(
                ok=False,
                mode="fast",
                kind="outline",
                message=(
                    f"목차 파일이 없습니다.\n"
                    f"- 기대 경로: {outline_path_p}\n\n"
                    f"먼저 (대화형) 목차 생성 흐름을 태우거나, 위 경로에 {outline_path_p.name} 파일을 만들어 주세요."
                ),
                last_saved_path=None,
                state_patch=patch,
                artifacts={"outline_path": str(outline_path_p), "topic_slug": slug},
            )

        outline_text = outline_path_p.read_text(encoding="utf-8", errors="replace")
        outline_text = self._truncate(outline_text, max_chars=12000)

        patch = {
            "execution_mode": "fast",
            "outline_path": str(outline_path_p),
            "last_saved_path": str(outline_path_p),
            "flags": {"outline_shown": True},
        }
        self._apply_patch(state, patch)

        return FastResponse(
            ok=True,
            mode="fast",
            kind="outline",
            message=outline_text,
            last_saved_path=str(outline_path_p),
            state_patch=patch,
            artifacts={"outline_path": str(outline_path_p), "topic_slug": slug},
        )

    def _handle_status(self, state: MutableMapping[str, Any]) -> FastResponse:
        slug = self._topic_slug(state) or "(unset)"
        pending = state.get("pending")
        if not isinstance(pending, list):
            pending = []
        last_saved = state.get("last_saved_path")
        mode = state.get("execution_mode") or "(unset)"
        raw_flags = state.get("flags")
        flags: Dict[str, Any] = raw_flags if isinstance(raw_flags, dict) else {}

        msg = (
            f"✅ STATUS\n"
            f"- topic_slug: {slug}\n"
            f"- execution_mode: {mode}\n"
            f"- pending: {pending}\n"
            f"- last_saved_path: {last_saved}\n"
            f"- flags(outline_shown): {flags.get('outline_shown')}\n"
        )
        patch = {"execution_mode": "fast"}
        self._apply_patch(state, patch)
        return FastResponse(ok=True, mode="fast", kind="status", message=msg, state_patch=patch)

    def _handle_help(self) -> FastResponse:
        msg = (
            "✅ 사용 가능한 Fast 명령\n"
            "- /outline : 목차 표시(LLM 0회)\n"
            "- /new-topic <slug> : 토픽 전환/초기화(LLM 0회)\n"
            "- /write <exact title> : 쓰기 예약만(LLM 0회)\n"
            "- /status : 상태 요약(LLM 0회)\n"
            "- /set KEY=VALUE : 설정 변경(LLM 0회)\n"
            "\n"
            "자연어 별칭(보수적): '목차', '목차 보여줘'\n"
        )
        return FastResponse(ok=True, mode="fast", kind="help", message=msg, state_patch={"execution_mode": "fast"})

    def _handle_new_topic(self, arg: str, state: MutableMapping[str, Any]) -> FastResponse:
        slug = (arg or "").strip()
        if not slug:
            patch = {"execution_mode": "fast"}
            self._apply_patch(state, patch)
            return FastResponse(
                ok=False, mode="fast", kind="new_topic",
                message="사용법: /new-topic <slug>",
                state_patch=patch
            )

        # slug normalize (deterministic)
        slug_norm = re.sub(r"[^a-zA-Z0-9-]+", "-", slug).strip("-").lower()
        if not slug_norm:
            patch = {"execution_mode": "fast"}
            self._apply_patch(state, patch)
            return FastResponse(
                ok=False, mode="fast", kind="new_topic",
                message="유효한 slug가 아닙니다. 영문/숫자/대시 조합을 권장합니다.",
                state_patch=patch
            )

        # Minimal topic reset (keep conservative: only clearly topic-bound keys)
        patch = {
            "execution_mode": "fast",
            "topic_slug": slug_norm,
            "flags": {"outline_shown": False},
            "pending": [],
            "refs": [],
        }
        self._apply_patch(state, patch)

        # ✅ core/paths.py 규칙으로 outline 경로 확정 후 디렉터리만 준비
        outline_path_p = self.DEFAULT_OUTLINE_PATH(state, slug_norm)
        outline_path_p.parent.mkdir(parents=True, exist_ok=True)

        return FastResponse(
            ok=True,
            mode="fast",
            kind="new_topic",
            message=(
                f"✅ 토픽 전환 완료: {slug_norm}\n"
                f"- outline 경로: {outline_path_p}\n"
                f"다음: /outline 또는 (대화형) '목차 생성' 입력"
            ),
            last_saved_path=None,
            state_patch=patch,
            artifacts={"topic_slug": slug_norm, "outline_path": str(outline_path_p)},
        )

    def _handle_set(self, arg: str, state: MutableMapping[str, Any]) -> FastResponse:
        m = self._RE_SET.match(arg or "")
        if not m:
            patch = {"execution_mode": "fast"}
            self._apply_patch(state, patch)
            return FastResponse(
                ok=False, mode="fast", kind="set",
                message="사용법: /set KEY=VALUE (예: /set DIRECT_QA=1)",
                state_patch=patch
            )

        key = m.group("key").strip()
        val_raw = m.group("val").strip()

        # allowlist (keep tight)
        ALLOW = {"DIRECT_QA", "AUTO_WRITE_AFTER_RAG", "AUTO_WRITE_DURING_RESEARCH", "DOC_MODE"}
        if key not in ALLOW:
            patch = {"execution_mode": "fast"}
            self._apply_patch(state, patch)
            return FastResponse(
                ok=False, mode="fast", kind="set",
                message=f"허용되지 않은 KEY: {key}\n허용: {sorted(ALLOW)}",
                state_patch=patch
            )

        # deterministic parse
        if key in {"DIRECT_QA", "AUTO_WRITE_AFTER_RAG", "AUTO_WRITE_DURING_RESEARCH"}:
            v = val_raw.lower()
            if v in ("1", "true", "yes", "y", "on"):
                val = True
            elif v in ("0", "false", "no", "n", "off"):
                val = False
            else:
                patch = {"execution_mode": "fast"}
                self._apply_patch(state, patch)
                return FastResponse(ok=False, mode="fast", kind="set",
                                    message=f"값 파싱 실패: {key}={val_raw} (0/1, true/false만 허용)",
                                    state_patch=patch)
            patch = {"execution_mode": "fast", "flags": {key: val}}
            self._apply_patch(state, patch)
            return FastResponse(ok=True, mode="fast", kind="set",
                                message=f"✅ {key}={1 if val else 0} 로 설정했습니다.",
                                state_patch=patch)

        if key == "DOC_MODE":
            v = val_raw.lower()
            if v not in ("report", "book"):
                patch = {"execution_mode": "fast"}
                self._apply_patch(state, patch)
                return FastResponse(ok=False, mode="fast", kind="set",
                                    message="DOC_MODE는 report 또는 book만 허용합니다.",
                                    state_patch=patch)
            patch = {"execution_mode": "fast", "flags": {"DOC_MODE": v}}
            self._apply_patch(state, patch)
            return FastResponse(ok=True, mode="fast", kind="set",
                                message=f"✅ DOC_MODE={v} 로 설정했습니다.",
                                state_patch=patch)

        # unreachable
        patch = {"execution_mode": "fast"}
        self._apply_patch(state, patch)
        return FastResponse(ok=False, mode="fast", kind="set", message="처리 불가", state_patch=patch)

    def _handle_write_queue(self, arg: str, state: MutableMapping[str, Any]) -> FastResponse:
        title = (arg or "").strip()
        if not title:
            patch = {"execution_mode": "fast"}
            self._apply_patch(state, patch)
            return FastResponse(ok=False, mode="fast", kind="write_queued",
                                message="사용법: /write <exact title>",
                                state_patch=patch)

        # 예약만. LLM/graph 금지.
        flags_patch = {"requested_write_title": title}
        patch = {"execution_mode": "write", "flags": flags_patch}

        # pending add (deterministic, no duplicates)
        pending = state.get("pending")
        if not isinstance(pending, list):
            pending = []
        if "section_writer" not in pending:
            pending.append("section_writer")
        state["pending"] = pending

        self._apply_patch(state, patch)

        return FastResponse(
            ok=True,
            mode="fast",
            kind="write_queued",
            message=f"✅ 쓰기 예약: {title}\n(다음 /run 호출에서 실행 흐름이 이어집니다)",
            last_saved_path=None,
            state_patch={**patch, "pending": pending},
            artifacts={"pending_added": ["section_writer"], "requested_write_title": title},
        )

def run_once(state: State, user_input: str, recursion_limit: int) -> State:
    """
    콘솔 루프에서 하던 '입력 1회 처리'를 함수화.
    - state에 HumanMessage 추가
    - 그래프 invoke
    - 결과 merge
    - save_state
    """

    if not str(user_input).strip():
        return state

    # --- FAST COMMAND ROUTER (LLM 0회 보장) ---
    fast = FastCommandHandler()

    norm = " ".join(str(user_input).split())

    # TypedDict(State)는 MutableMapping으로 간주되지 않는 타입체커가 있어
    # Fast는 임시 dict를 mutate하고, 결과를 State에 반영한다.
    fast_state: dict[str, Any] = dict(cast(dict[str, Any], state))
    handled, fresp = fast.handle(user_input, fast_state)

    logger.info(
        "[RUN_ONCE_FAST] input=%r handled=%s kind=%s ok=%s",
        user_input,
        handled,
        getattr(fresp, "kind", None),
        getattr(fresp, "ok", None),
    )

    if handled and fresp is not None:
        state.setdefault("messages", []).append(
            HumanMessage(content=user_input, additional_kwargs={"fast_path": True})
        )
        # ✅ 여기서 graph/supervisor 호출하면 안 됨
        # TypedDict(State)는 update(dict[str, Any])를 타입상 허용하지 않음(pyright).
        # Fast가 변경할 수 있는 키만 allowlist로 안전하게 반영한다.
        _FAST_KEYS: tuple[str, ...] = (
            "execution_mode",
            "outline_path",
            "last_saved_path",
            "topic_slug",
            "pending",
            "refs",
            "flags",
        )
        st_any = cast(dict[str, Any], state)
        for k in _FAST_KEYS:
            if k in fast_state:
                if k == "flags" and isinstance(fast_state[k], dict):
                    # merge flags (do not clobber)
                    existing = st_any.get("flags")
                    if not isinstance(existing, dict):
                        existing = {}
                    existing.update(fast_state[k])
                    st_any["flags"] = existing
                else:
                    st_any[k] = fast_state[k]
        # Fast 응답을 "State"에 남겨 API/communicator가 그대로 전달할 수 있게 한다.
        # (run_once는 State를 반환해야 함)
        if fresp.last_saved_path:
            state["last_saved_path"] = fresp.last_saved_path

        # fast 응답은 LLM이 만든 게 아니므로 qa_direct_reply는 False 유지
        state["qa_direct_reply"] = False
        state.setdefault("messages", []).append(
            AIMessage(
                content=fresp.message,
                additional_kwargs={
                    "fast_path": True,
                    "fast_kind": fresp.kind,
                    "fast_ok": fresp.ok,
                    "state_patch": fresp.state_patch or {},
                    "artifacts": fresp.artifacts or {},
                },
            )
        )

        # 저장 후 즉시 리턴 (그래프 invoke 금지)
        try:
            save_state(current_path, state)
        except Exception:
            logger.exception("save_state failed (non-fatal)")
        return state


    # build report 커맨드(콘솔에서만 쓰던 특수 명령도 웹에서 그대로 동작시키고 싶으면 여기에 둡니다)
    _u = user_input.strip().lower()
    if _u in ("build: report", "report build", "보고서 빌드", "최종 보고서 생성"):
        slug = str(state.get("topic_slug") or "default")
        outline_fname = state.get("outline_fname") or ("outline_report.md" if config.DOC_MODE == "report" else "outline.md")
        out_path, missing = build_final_report(
            topic_slug=slug,
            outline_fname=outline_fname,
            mode=config.DOC_MODE,
            root_dir=str(current_path)
        )
        update_flags(state)
        state["last_saved_path"] = out_path
        # 누락 섹션은 state에 남겨두고, 텍스트는 응답에서 사용
        return state
    
    # ✅ CommandRouter: 이번 입력의 “명령 의도”를 1회만 파싱해 flags에 저장
    try:
        flags = dict(state.get("flags") or {})
        flags["command_intent"] = parse_command_intent(user_input)
        state["flags"] = flags
    except Exception:
        logger.exception("parse_command_intent failed (non-fatal)")

    # fast-hit이 아닌 경우에만 user message를 그래프 입력으로 추가
    state.setdefault("messages", []).append(HumanMessage(content=user_input))

    # (기존 show_outline short-circuit 블록은 PRE-FLIGHT로 대체됨)

    # ✅ intent 기반 선주입: create_outline만 남긴다.
    # show_outline은 위에서 이미 short-circuit(return) 처리하므로 여기서 또 예약할 필요가 없음.
    try:
        intent = ((state.get("flags") or {}).get("command_intent") or {})
        itype = str(intent.get("type") or "")
        if itype == "create_outline":
            # outline 파일명은 state 기준으로 안전하게 결정
            fname = str(state.get("outline_fname") or ("outline_report.md" if config.DOC_MODE == "report" else "outline.md"))
            tasks = list(state.get("task_history") or [])
            if not _has_pending_task(state, "content_strategist"):
                tasks.append(Task(agent="content_strategist", done=False, description=f"create_outline:{fname}", done_at=""))
            state["qa_direct_reply"] = False
            # ✅ TypedDict 안전: task_history는 로컬 리스트로 조작 후 되돌려 넣는다
            state["task_history"] = tasks
    except Exception:
        logger.exception("intent pre-inject failed (non-fatal)")

    # 그래프 획득(필요 시 재빌드)
    graph_obj = _get_graph_cached()

    def _invoke_graph(g: Any, st: Any, cfg: dict | None = None) -> Any:
        if hasattr(g, "invoke"):
            return g.invoke(st, config=cfg)
        if hasattr(g, "run"):
            return g.run(st, config=cfg)
        raise TypeError("Graph object exposes neither 'invoke' nor 'run'.")

    result = _invoke_graph(graph_obj, state, {"recursion_limit": recursion_limit})

    if not isinstance(result, dict):
        raise TypeError(f"graph.invoke returned unexpected type: {type(result).__name__}")

    merged = dict(state)
    merged.update(result)
    state = cast(State, merged)

    # 저장
    try:
        save_state(current_path, state)
    except Exception as e:
        logger.exception("save_state failed: %s", e)

    return state

@web_app.get("/api/health")
def health():
    return {"ok": True}


@web_app.get("/api/state")
def api_state():
    global _WEB_STATE
    st = _WEB_STATE
    if st is None:
        return {"ok": False, "error": "state_not_initialized"}
    
    phase = "running" if _RUN_LOCK.locked() else "idle"
    iter_count = int(st.get("iteration_count") or 0)

    try:
        update_flags(st, iteration_count=iter_count)
    except Exception:
        pass

    # ── 토픽 미션 (BLOCKAGI_OBJECTIVE_1/2/3) 수집 ──
    # 비어있지 않은 것만 배열에 담음. 환경변수가 없으면 빈 배열.
    objectives = []
    for i in range(1, 11):  # 1~10까지 여유있게 스캔
        val = os.environ.get(f"BLOCKAGI_OBJECTIVE_{i}", "").strip()
        if val:
            objectives.append(val)

    return {
        "ok": True,
        "doc_mode": getattr(config.CFG, "DOC_MODE", None),
        "namespace": (st.get("flags") or {}).get("chroma_namespace") or None,
        "pending": len(st.get("task_history", []) or []),
        "refs": len(((st.get("references") or {}).get("docs") or [])),
        "last_saved_path": st.get("last_saved_path"),
        "outline": None,
        "flags": st.get("flags") or {},
        "objectives": objectives,   # ← 새 필드 추가
        "phase": phase,
        "cancel_requested": bool(_CANCEL_EVENT.is_set()),
        "iteration_count": iter_count,
        "updated_at": _now_str(),
    }

@web_app.post("/api/cancel")
def api_cancel():
    """
    실행 취소 요청 플래그를 설정합니다.
    - 주의: run_once 내부가 이 플래그를 체크하지 않으면 즉시 중단되진 않습니다.
    """
    _CANCEL_EVENT.set()
    return {"ok": True}


@web_app.get("/api/outline")
def api_outline():
    """
    return:
      {"ok": True, "items": [...], "path": "...", "source": "topic|default|empty"}
    """
    global _WEB_STATE
    st = _WEB_STATE
    if st is None:
        return {"ok": False, "error": "state_not_initialized"}

    p_topic, p_default = _outline_paths(st)

    if os.path.exists(p_topic):
        path = p_topic
        source = "topic"
    elif os.path.exists(p_default):
        path = p_default
        source = "default"
    else:
        return {"ok": True, "items": [], "path": p_topic, "source": "empty"}

    with open(path, "r", encoding="utf-8") as f:
        lines = [ln.rstrip("\n") for ln in f.readlines() if ln.strip()]

    return {"ok": True, "items": lines, "path": path, "source": source}


@web_app.put("/api/outline")
def api_outline_save(payload: Dict[str, Any]):
    """
    payload:
      {"items": ["1. 서론", "2. ...", ...]}
    저장:
      outlines/<slug>/<outline_fname>
    """
    global _WEB_STATE
    st = _WEB_STATE
    if st is None:
        return {"ok": False, "error": "state_not_initialized"}

    items = payload.get("items")
    if not isinstance(items, list):
        return {"ok": False, "error": "items_must_be_list"}

    p_topic, _p_default = _outline_paths(st)
    os.makedirs(os.path.dirname(p_topic), exist_ok=True)

    cleaned: list[str] = []
    for it in items:
        s = str(it).strip()
        if s:
            cleaned.append(s)

    with open(p_topic, "w", encoding="utf-8") as f:
        for s in cleaned:
            f.write(s + "\n")

    return {"ok": True, "path": p_topic, "count": len(cleaned)}



@web_app.post("/api/run")
async def api_run(payload: Dict[str, Any]):
    """
    Next.js에서 보내는 형태(권장):
      {"input":"목차 보여줘", "options": {...}}
    """
    global _WEB_STATE, _WEB_ARGS

    if _WEB_STATE is None:
        return {"ok": False, "error": "state_not_initialized"}

    user_input = str(payload.get("input") or "").strip()
    print("[API_RUN_PRINT] input=", repr(user_input), flush=True)  # ✅ 여기
    if not user_input:
        return {"ok": False, "error": "empty_input"}

    # 새 실행 시작 시 cancel 플래그 해제
    _CANCEL_EVENT.clear()

    # 새 실행 시작 시 이벤트 버퍼 초기화 — 사용자 관점에서는 명령 한 번이 한 흐름
    clear_events()
    emit_event("작업 시작", kind="start", detail=user_input[:120])


    # 옵션은 우선 받아두기만 하고(최소 수정),
    # 실제로 CFG/ENV 반영까지 하려면 2단계에서 확장하면 됩니다.
    # options = payload.get("options") or {}

    recursion_limit = int(getattr(_WEB_ARGS, "recursion_limit", 200) or 200)

    async with _RUN_LOCK:
        try:
            # CPU/IO가 섞인 무거운 작업이므로 executor에서 돌리는 게 안전
            loop = asyncio.get_running_loop()
            _WEB_STATE = await loop.run_in_executor(None, run_once, _WEB_STATE, user_input, recursion_limit)
            state_now = _WEB_STATE
            assert state_now is not None  # ✅ Pylance에게 "None 아님" 확정

            emit_event("작업 완료", kind="done")
            return {
                "ok": True,
                "message": _last_ai_text(state_now) or "OK",
                "last_saved_path": state_now.get("last_saved_path"),
            }
        except Exception as e:
            logger.exception("api_run failed: %s", e)
            emit_event("오류 발생", kind="error", detail=str(e)[:200])
            return {"ok": False, "error": str(e)}
        
# ═══════════════════════════════════════════════════════════
# Export 엔드포인트 — 마크다운 → docx 변환
# ═══════════════════════════════════════════════════════════

class ExportRequest(BaseModel):
    kind: str  # "section" | "report"
    section_id: Optional[int] = None
    format: str = "docx"


def _load_outline_items() -> list[str]:
    """현재 활성 outline의 줄 단위 항목 목록 반환. 실패 시 빈 리스트."""
    global _WEB_STATE
    st = _WEB_STATE
    if st is None:
        return []
    try:
        p_topic, p_default = _outline_paths(st)
        path = p_topic if os.path.exists(p_topic) else (
            p_default if os.path.exists(p_default) else None
        )
        if not path:
            return []
        with open(path, "r", encoding="utf-8") as f:
            return [ln.rstrip("\n") for ln in f.readlines() if ln.strip()]
    except Exception:
        return []


def _outline_title_to_slug(raw: str) -> str:
    """outline 한 줄에서 제목만 추출 후 슬러그화 (frontend lib/data.ts:slugifyTitle와 동치)."""
    from utils.text_utils import slugify
    s = re.sub(r"^\s*#+\s*", "", raw)
    s = re.sub(r"^\s*\d+[.)]\s*", "", s).strip()
    return slugify(s, allow_unicode=True)


def _resolve_section_file(
    section_id: int, slug_dir: str, outline_items: list[str]
) -> str | None:
    """section_id → .md 절대 경로. 못 찾으면 None.
    1) 파일명 prefix "{section_id}-..." (옛 형식)
    2) outline_items[section_id-1] 제목 슬러그가 파일 stem과 일치 (신 형식 — writer가 prefix 안 붙임)
    """
    files = [f for f in os.listdir(slug_dir) if f.endswith(".md")]

    for fname in files:
        if fname.startswith(f"{section_id}-"):
            return os.path.join(slug_dir, fname)

    if 1 <= section_id <= len(outline_items):
        target = _outline_title_to_slug(outline_items[section_id - 1]).lower()
        if target:
            for fname in files:
                stem = re.sub(r"\.md$", "", fname).lower()
                if stem == target:
                    return os.path.join(slug_dir, fname)

    return None


def _read_section_file(section_id: int) -> tuple[str, str]:
    """sections/<slug>/ 의 섹션 파일을 읽어서 (제목, 본문) 반환."""
    sections_dir = os.path.join(str(current_path), "sections")
    if not os.path.isdir(sections_dir):
        raise HTTPException(404, "sections directory not found")

    slug_dirs = [
        d for d in os.listdir(sections_dir)
        if os.path.isdir(os.path.join(sections_dir, d))
    ]
    if not slug_dirs:
        raise HTTPException(404, "no slug directory")

    # 활성 slug — 가장 최근 수정된 폴더를 기본으로
    target_slug = sorted(
        slug_dirs,
        key=lambda d: os.path.getmtime(os.path.join(sections_dir, d)),
        reverse=True,
    )[0]
    slug_dir = os.path.join(sections_dir, target_slug)

    outline_items = _load_outline_items()
    fpath = _resolve_section_file(section_id, slug_dir, outline_items)
    if not fpath:
        raise HTTPException(404, f"section {section_id} not found")

    with open(fpath, "r", encoding="utf-8") as f:
        content = f.read()
    title_match = re.search(r"^##\s+(.+)$", content, re.MULTILINE)
    if title_match:
        title = title_match.group(1).strip()
        title = re.sub(r"^\d+[.)]\s*", "", title)
    else:
        title = os.path.basename(fpath)
    return title, content


def _read_all_sections() -> list[tuple[int, str, str]]:
    """모든 섹션을 (id, title, content) 리스트로 반환 (번호 순 정렬).
    매칭: 파일명 prefix 숫자 우선, 없으면 outline 제목 슬러그 역매핑.
    """
    sections_dir = os.path.join(str(current_path), "sections")
    if not os.path.isdir(sections_dir):
        return []

    slug_dirs = [
        d for d in os.listdir(sections_dir)
        if os.path.isdir(os.path.join(sections_dir, d))
    ]
    if not slug_dirs:
        return []

    target_slug = sorted(
        slug_dirs,
        key=lambda d: os.path.getmtime(os.path.join(sections_dir, d)),
        reverse=True,
    )[0]
    slug_dir = os.path.join(sections_dir, target_slug)

    outline_items = _load_outline_items()
    slug_to_id: dict[str, int] = {}
    for i, raw in enumerate(outline_items):
        sl = _outline_title_to_slug(raw).lower()
        if sl:
            slug_to_id[sl] = i + 1

    results: list[tuple[int, str, str]] = []
    for fname in os.listdir(slug_dir):
        if not fname.endswith(".md"):
            continue
        sid: int | None = None
        m = re.match(r"^(\d+)-", fname)
        if m:
            sid = int(m.group(1))
        else:
            stem = re.sub(r"\.md$", "", fname).lower()
            sid = slug_to_id.get(stem)
        if sid is None:
            continue
        fpath = os.path.join(slug_dir, fname)
        with open(fpath, "r", encoding="utf-8") as f:
            content = f.read()
        title_match = re.search(r"^##\s+(.+)$", content, re.MULTILINE)
        if title_match:
            title = title_match.group(1).strip()
            title = re.sub(r"^\d+[.)]\s*", "", title)
        else:
            title = fname
        results.append((sid, title, content))

    results.sort(key=lambda x: x[0])
    return results


def _markdown_to_docx(blocks: list[tuple[int, str, str]], doc_title: str) -> bytes:
    """
    여러 섹션을 받아 하나의 docx로 변환.
    blocks: [(section_id, section_title, markdown_content), ...]
    반환: docx 바이트
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # 기본 폰트 설정 (한글 호환)
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"   # type: ignore[attr-defined]
    style.font.size = Pt(11)    # type: ignore[attr-defined]

    # 문서 제목
    title_para = doc.add_heading(doc_title, level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    for sid, stitle, content in blocks:
        # 섹션 제목
        doc.add_heading(f"{sid}. {stitle}", level=1)

        # 본문/footnote 분리
        body_text, footnotes = _split_body_footnotes(content)

        # 본문 렌더링
        _render_markdown_to_docx(doc, body_text)

        # footnote 영역
        if footnotes:
            doc.add_paragraph()
            fn_heading = doc.add_heading("참고 문헌", level=3)
            for fn_marker, fn_text in footnotes:
                p = doc.add_paragraph()
                p.add_run(f"[{fn_marker}] ").bold = True
                p.add_run(fn_text)

    # 바이트로 반환
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _split_body_footnotes(content: str) -> tuple[str, list[tuple[str, str]]]:
    """
    마크다운에서 본문과 footnote 분리.
    반환: (본문 텍스트, [(marker, decoded_text), ...])
    """
    lines = content.split("\n")
    fn_start = -1
    for i, line in enumerate(lines):
        if line.strip() == "---":
            for j in range(i + 1, len(lines)):
                next_line = lines[j].strip()
                if not next_line:
                    continue
                if re.match(r"^#{2,4}\s*(참고\s*문헌|각주)", next_line):
                    fn_start = i
                break
            if fn_start != -1:
                break

    if fn_start == -1:
        return content, []

    body = "\n".join(lines[:fn_start])
    fn_lines = lines[fn_start + 1:]

    footnotes = []
    for line in fn_lines:
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        m = re.match(r"^\[\^([^\]]+)\]:\s*(.+)$", line)
        if m:
            marker = m.group(1)
            text = m.group(2).strip()
            # URL 디코딩 (file://%EC%95%84... → 한글)
            try:
                from urllib.parse import unquote
                # 라벨 뒤에 (...) 부분만 추출하면 깔끔
                label_match = re.match(r"^.+?\s+\((.+)\)\s*$", text)
                if label_match:
                    text = label_match.group(1).strip()
                else:
                    text = unquote(text)
            except Exception:
                pass
            footnotes.append((marker, text))

    return body, footnotes


def _render_markdown_to_docx(doc, markdown: str):
    """
    간이 마크다운 → docx 렌더링.
    지원: ## 헤딩, ### 헤딩, **bold**, [파일명] 인용 (파란색), 1. 리스트, - 리스트
    """
    from docx.shared import Pt, RGBColor

    lines = markdown.split("\n")
    paragraph_buffer = []

    def flush_paragraph():
        if not paragraph_buffer:
            return
        text = " ".join(paragraph_buffer)
        p = doc.add_paragraph()
        _add_inline_runs(p, text)
        paragraph_buffer.clear()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            continue

        # 첫 헤딩(## N. ...)은 이미 추가했으니 스킵
        if re.match(r"^##\s+\d+\.\s+", stripped):
            continue

        if stripped == "---":
            flush_paragraph()
            continue

        # 헤딩
        h_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if h_match:
            flush_paragraph()
            level = min(len(h_match.group(1)) + 1, 4)
            doc.add_heading(h_match.group(2).strip(), level=level)
            continue

        # 순서있는 리스트
        ol_match = re.match(r"^\s*(\d+)[.)]\s+(.+)$", line)
        if ol_match:
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, ol_match.group(2))
            continue

        # 순서없는 리스트
        ul_match = re.match(r"^\s*[-*]\s+(.+)$", line)
        if ul_match:
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, ul_match.group(1))
            continue

        paragraph_buffer.append(stripped)

    flush_paragraph()


def _add_inline_runs(paragraph, text: str):
    """
    단락 내 인라인 마크업 처리.
    **bold** → 굵게
    [파일명.확장자] → 파란색 칩처럼
    그 외 → 일반 텍스트
    """
    from docx.shared import RGBColor

    pattern = re.compile(
        r"(\*\*[^*]+\*\*)|(\[[^\]\n]*\.[a-zA-Z0-9가-힣_-]+[^\]\n]*\])"
    )
    last_idx = 0
    for m in pattern.finditer(text):
        if m.start() > last_idx:
            paragraph.add_run(text[last_idx:m.start()])

        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("["):
            inner = token[1:-1]
            run = paragraph.add_run(f"[{inner}]")
            run.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5)  # 인용 칩 파란색
            run.font.size = None  # 그대로
        last_idx = m.end()

    if last_idx < len(text):
        paragraph.add_run(text[last_idx:])


@web_app.post("/api/export")
def api_export(req: ExportRequest):
    """
    Word(.docx) 다운로드.
    payload: { kind: "section"|"report", section_id?: int, format: "docx" }
    """
    if req.format != "docx":
        raise HTTPException(400, f"unsupported format: {req.format}")

    topic_title = os.environ.get("TOPIC_TITLE", "보고서")

    if req.kind == "section":
        if req.section_id is None:
            raise HTTPException(400, "section_id required for kind=section")
        title, content = _read_section_file(req.section_id)
        blocks = [(req.section_id, title, content)]
        doc_title = f"{topic_title} — {req.section_id}. {title}"
        filename = f"{req.section_id}-{_safe_filename(title)}.docx"
    elif req.kind == "report":
        all_sections = _read_all_sections()
        if not all_sections:
            raise HTTPException(404, "no sections found")
        blocks = all_sections
        doc_title = topic_title
        filename = f"{_safe_filename(topic_title)}.docx"
    else:
        raise HTTPException(400, f"invalid kind: {req.kind}")

    docx_bytes = _markdown_to_docx(blocks, doc_title)

    from urllib.parse import quote
    
    # 한글 파일명을 위한 RFC 5987 인코딩
    # ASCII fallback (filename=) + UTF-8 (filename*=) 둘 다 제공
    encoded_filename = quote(filename, safe="")
    ascii_fallback = filename.encode("ascii", errors="ignore").decode("ascii") or "report.docx"
    
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": (
                f'attachment; filename="{ascii_fallback}"; '
                f"filename*=UTF-8''{encoded_filename}"
            ),
        },
    )


def _safe_filename(s: str) -> str:
    """파일명에 부적합한 문자 제거"""
    s = re.sub(r'[\\/:*?"<>|]', "_", s)
    s = s.strip().replace(" ", "_")
    return s[:80]

@web_app.get("/api/files")
def api_files(kind: str | None = None, limit: int = 200):
    """
    중요한 산출물 목록 제공.

    기본 동작:
      - doc_mode=report => sections/<slug>/*.md
      - doc_mode=book   => chapters/<slug>/*.md

    query:
      - kind: "artifact"(기본) | "report"(=reports) | "reportlog"(호환)
      - limit: 최대 반환(기본 200, 최대 500)

    return:
      {"ok": True, "mode": "...", "slug":"...", "root":"sections|chapters",
       "files":[{"id":"...", "name":"...", "path":"...", "mtime":..., "size":...}]}
    """
    global _WEB_STATE
    st = _WEB_STATE
    if st is None:
        return {"ok": False, "error": "state_not_initialized"}

    slug = _topic_slug_from_state(st)
    mode = _doc_mode_from_cfg()
    kind_norm = (kind or "artifact").strip().lower()
    if kind_norm in ("report", "reports", "reportlog"):
        root = "reports"
        adir = os.path.join(str(current_path), "reports", slug)
    else:
        root = _artifact_root_for_mode(mode)
        adir = _artifact_dir(mode, slug)

    limit = max(1, min(int(limit or 200), 500))

    if not os.path.isdir(adir):
        return {"ok": True, "mode": mode, "slug": slug, "root": root, "files": []}

    files = []
    for name in os.listdir(adir):
        if not name.lower().endswith(".md"):
            continue
        full = os.path.join(adir, name)
        if not os.path.isfile(full):
            continue

        try:
            st_m = os.stat(full)
            mtime = int(st_m.st_mtime)
            size = int(st_m.st_size)
        except Exception:
            mtime = 0
            size = 0

        # id 형식: sections|chapters|reports/<slug>/<name>
        file_id = f"{root}/{slug}/{name}"
        files.append({"id": file_id, "name": name, "path": full, "mtime": mtime, "size": size})

    files.sort(key=lambda x: x.get("mtime", 0), reverse=True)
    return {"ok": True, "mode": mode, "slug": slug, "root": root, "files": files[:limit]}


@web_app.get("/api/files/{file_id:path}")
def api_file_download(file_id: str):
    """
    file_id 예:
      sections/<slug>/<file>.md
      chapters/<slug>/<file>.md
      reports/<slug>/<file>.md

    sections/ 또는 chapters/ 또는 reports/ 아래만 허용.
    """
    target = os.path.join(str(current_path), file_id)

    if not _safe_under_artifacts(target):
        return {"ok": False, "error": "forbidden_path"}

    if not os.path.exists(target) or not os.path.isfile(target):
        return {"ok": False, "error": "not_found"}

    # 마크다운/텍스트 파일은 JSON으로 감싸서 반환 (브라우저 fetch 호환)
    try:
        with open(target, "r", encoding="utf-8") as f:
            content = f.read()
        return {
            "ok": True,
            "id": file_id,
            "name": os.path.basename(target),
            "content": content,
            "size": os.path.getsize(target),
        }
    except UnicodeDecodeError:
        # 텍스트로 읽을 수 없는 바이너리 파일은 기존대로 FileResponse
        return FileResponse(
            target,
            media_type="application/octet-stream",
            filename=os.path.basename(target),
        )


@web_app.get("/api/section-refs/{file_id:path}")
def api_section_refs(file_id: str):
    """
    섹션 .md 옆의 사이드카 .refs.json 을 반환.
    file_id 예: sections/<slug>/<file>.md  (확장자 .md 무관 — 자동으로 .refs.json 매핑)

    응답:
      {"ok": True, "id": <file_id>, "refs": {"1": {marker, url, label, text, source, title}, ...}}
      파일 없으면 {"ok": True, "id": <file_id>, "refs": {}} (404 대신 빈 맵 — 프런트 단순화)
    """
    import json as _json
    target_md = os.path.join(str(current_path), file_id)

    if not _safe_under_artifacts(target_md):
        return {"ok": False, "error": "forbidden_path"}

    # .md → .refs.json (다른 확장자가 들어와도 .refs.json 으로 일괄 치환)
    base, _ext = os.path.splitext(target_md)
    sidecar_path = base + ".refs.json"

    if not os.path.exists(sidecar_path) or not os.path.isfile(sidecar_path):
        return {"ok": True, "id": file_id, "refs": {}}

    try:
        with open(sidecar_path, "r", encoding="utf-8") as f:
            data = _json.load(f)
        if not isinstance(data, dict):
            data = {}
        return {"ok": True, "id": file_id, "refs": data}
    except Exception as _e:
        return {"ok": False, "error": f"read_failed: {_e}"}


@web_app.get("/api/logs")
def api_logs(cursor: int = 0, limit: int = 200):
    """
    폴링 방식 로그.
    - cursor: 마지막으로 받은 seq (없으면 0)
    - limit: 최대 반환 줄 수
    return:
      {"ok": True, "next_cursor": <int>, "lines": [{"seq":..., "line":"..."}]}
    """
    limit = max(1, min(int(limit or 200), 500))

    with _LOG_LOCK:
        # cursor보다 큰 것만
        out = [{"seq": seq, "line": line} for (seq, line) in _LOG_BUF if seq > int(cursor)]
        if len(out) > limit:
            out = out[-limit:]
        next_cursor = out[-1]["seq"] if out else int(cursor)

    return {"ok": True, "next_cursor": next_cursor, "lines": out}


@web_app.get("/api/events")
def api_events(cursor: int = 0, limit: int = 200):
    """
    사용자 관점 진행 이벤트 폴링.
    - cursor: 마지막으로 받은 seq (없으면 0)
    - limit: 최대 반환 줄 수
    return:
      {"ok": True, "next_cursor": <int>,
       "events": [{"seq":..., "ts":..., "label":"...", "kind":"...", "detail":"..."}]}
    """
    next_cursor, events = get_events_since(int(cursor or 0), int(limit or 200))
    return {"ok": True, "next_cursor": next_cursor, "events": events}


# ── 엔트리포인트 ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    # 0) 반드시 가장 먼저: ENV → CFG 반영 (초기 바인딩 깨끗하게)
    try:
        import core.config as _early_config
        _early_config.reload_config()
    except Exception as _e:
        print(f"[WARN] reload_config() failed: {_e}")

    parser = argparse.ArgumentParser()
    parser.add_argument("--iteration_count", type=str,
                        default=str(getattr(config.CFG, "ITERATION_COUNT", 3)))
    parser.add_argument("--agent_role", type=str,
                        default=(getattr(config.CFG, "AGENT_ROLE", "") or "").strip().lower())
    parser.add_argument("--recursion_limit", type=int, default=int(os.getenv("RECURSION_LIMIT", "200")))
    parser.add_argument("--log-level", type=str, default=getattr(config.CFG, "LOG_LEVEL", "INFO"))
    parser.add_argument("--log-file", type=str, default=os.getenv("LOG_FILE"))
    parser.add_argument("--log-json", action="store_true", default=bool(getattr(config.CFG, "LOG_JSON", False)))
    parser.add_argument("--topic-slug", type=str, default=getattr(config.CFG, "TOPIC_SLUG", "") or "default")

    parser.add_argument("--log-topk", type=int, default=_int_env("LOG_TOPK", 3))
    parser.add_argument("--log-dashboard", action="store_true",
                        default=_bool_env("LOG_DASHBOARD", True))
    parser.add_argument("--log-wrap", type=int, default=_int_env("LOG_WRAP", 88))
    parser.add_argument("--human-logs-verbose", action="store_true")
    parser.add_argument("--human-logs", action="store_true")
    parser.add_argument("--echo-outline", action="store_true")
    parser.add_argument("--echo-sections", action="store_true")
    parser.add_argument("--echo-report", action="store_true")

    parser.add_argument("--gatekeep", action="store_true", dest="gatekeep")
    parser.add_argument("--no-gatekeep", action="store_false", dest="gatekeep")
    parser.set_defaults(gatekeep=None)

    parser.add_argument("--serve", action="store_true", help="Run as web server (FastAPI)")
    parser.add_argument("--host", type=str, default=os.getenv("HOST", "127.0.0.1"))
    parser.add_argument("--port", type=int, default=int(os.getenv("PORT", "8000")))

    parser.add_argument("--allow-domains", type=str, default=os.getenv("ALLOWED_DOMAINS", ""))
    parser.add_argument("--allow-subdomains", action="store_true", dest="allow_subdomains")
    parser.add_argument("--no-allow-subdomains", action="store_false", dest="allow_subdomains")
    parser.set_defaults(allow_subdomains=None)

    args = parser.parse_args()

    logger.info("ARGS topic_slug=%r | ENV.TOPIC_TITLE=%r",
                args.topic_slug, os.getenv("TOPIC_TITLE"))

    # 인자 → ENV 반영
    if args.log_level:
        os.environ["LOG_LEVEL"] = args.log_level.upper()
    if args.log_file:
        os.environ["LOG_FILE"] = args.log_file
    os.environ["LOG_JSON"] = "1" if args.log_json else "0"

    os.environ["LOG_TOPK"] = str(args.log_topk)
    os.environ["LOG_DASHBOARD"] = "1" if args.log_dashboard else "0"
    os.environ["LOG_WRAP"] = str(args.log_wrap)
    os.environ["DASH_SIMPLE"]   = os.getenv("DASH_SIMPLE", "1")
    os.environ["DASH_RATE_SEC"] = os.getenv("DASH_RATE_SEC", str(getattr(config.CFG, "DASH_RATE_SEC", 6)))
    os.environ["COMMUNICATOR_ECHO"] = os.getenv("COMMUNICATOR_ECHO", "0")
    os.environ["HUMAN_LOGS_VERBOSE"] = "1" if args.human_logs_verbose else ("1" if getattr(config.CFG, "HUMAN_LOGS_VERBOSE", False) else "0")
    os.environ["HUMAN_LOGS"]    = "1" if args.human_logs     else ("1" if getattr(config.CFG, "HUMAN_LOGS", False) else "0")
    os.environ["ECHO_OUTLINE"]  = "1" if args.echo_outline   else ("1" if getattr(config.CFG, "ECHO_OUTLINE", False) else "0")
    os.environ["ECHO_SECTIONS"] = "1" if args.echo_sections  else ("1" if getattr(config.CFG, "ECHO_SECTIONS", False) else "0")
    os.environ["ECHO_REPORT"]   = "1" if args.echo_report    else os.getenv("ECHO_REPORT", "0")

    if args.gatekeep is not None:
        os.environ["GATE_KEEP_SOURCES"] = "1" if args.gatekeep else "0"
    else:
        os.environ["GATE_KEEP_SOURCES"] = os.getenv("GATE_KEEP_SOURCES", "0")

    if args.allow_domains is not None and args.allow_domains.strip():
        os.environ["ALLOWED_DOMAINS"] = args.allow_domains

    if args.allow_subdomains is not None:
        os.environ["ALLOW_SUBDOMAINS"] = "1" if args.allow_subdomains else "0"

    # 인자→ENV 반영 후 CFG **in-place** 리로드(모듈 속성에 갱신하여 전역 일관 유지)
    try:
        config.CFG = config.reload_config()
    except Exception:
        logger.exception("reload_config() failed; continuing with previous CFG")
    # 게이트키핑 캐시 리프레시
    try:
        from settings_gatekeep import refresh_gatekeep_cache
        refresh_gatekeep_cache()
    except Exception:
        pass

    # 로깅 설정 (CFG 기반) — 여기서 파일 경로 확정 → faulthandler 파일 예약
    try:
        setup_logging()
        attach_web_log_handler()
    except Exception as e:
        # 로깅 설정 실패해도 덤프 타이머는 정리
        _cancel_fault_timers_and_close()
        raise

    # ▼ early debug buffer flush (DEBUG 모드에서만 존재)
    try:
        if _EARLY_DEBUG:
            _initlog = logging.getLogger("app.init")
            for _m in _EARLY_DEBUG:
                _initlog.debug(_m)
            _EARLY_DEBUG.clear()
    except Exception:
        pass
    # ▲

    logger.info(
        "ENV TOPIC_TITLE=%r | TOPIC_SLUG=%r | dotenv_path=%s",
        os.getenv("TOPIC_TITLE"),
        os.getenv("TOPIC_SLUG"),
        dotenv_path
    )

    try:
        from settings_gatekeep import gatekeep_enabled, get_allowed_domains
        if gatekeep_enabled():
            allow = ", ".join(sorted(get_allowed_domains())) or "(empty)"
            logger.info("[GATEKEEP] enabled; allowed=%s", allow)
        else:
            logger.info("[GATEKEEP] disabled")
    except Exception:
        pass

    iter_count = coerce_int(args.iteration_count, default=3)
    # role 우선순위: CLI → CFG.AGENT_ROLE → ENV.BLOCKAGI_AGENT_ROLE
    effective_role = (
        str(
            args.agent_role
            or getattr(config.CFG, "AGENT_ROLE", "")
            or os.getenv("BLOCKAGI_AGENT_ROLE", "")
        ).strip().lower()
        or None
    )


    # state에 반영하되, ENV는 최소한으로만 사용(호환 위해 TOPIC_* 유지)
    os.environ["TOPIC_SLUG"] = args.topic_slug
    if not os.getenv("TOPIC_TITLE"):
        os.environ["TOPIC_TITLE"] = args.topic_slug.replace("-", " ")

    state: State = initial_state(iteration_count=iter_count, agent_role=effective_role)

    # 웹 서버 모드에서 사용할 전역 상태 연결
    _WEB_STATE = state
    _WEB_ARGS = args


    # ── 초기 상태 시드: agent_role / research_objectives ──────────────────────
    try:
        from typing import MutableMapping, Any, List, cast
        mm = cast(MutableMapping[str, Any], state)
        # agent_role이 비어있다면 CFG.AGENT_ROLE로 보강
        if not (str(mm.get("agent_role") or "").strip()):
            mm["agent_role"] = (getattr(config.CFG, "AGENT_ROLE", "") or "").strip().lower() or None
        # 연구 목적: CFG.RESEARCH_OBJECTIVES → ENV 번호키 폴백
        if not mm.get("research_objectives"):
            objs: List[str] = []
            # 1) CFG에 있으면 그대로 사용
            try:
                _cfg_objs = list(getattr(config.CFG, "RESEARCH_OBJECTIVES", []) or [])
                objs.extend([s for s in _cfg_objs if isinstance(s, str) and s.strip()])
            except Exception:
                pass
            # 2) 비어있으면 ENV(BLOCKAGI_OBJECTIVE_1..N) 폴백
            if not objs:
                for k in ("BLOCKAGI_OBJECTIVE_1","BLOCKAGI_OBJECTIVE_2","BLOCKAGI_OBJECTIVE_3","BLOCKAGI_OBJECTIVE_4"):
                    v = os.getenv(k)
                    if v and v.strip():
                        objs.append(v.strip())
            if objs:
                mm["research_objectives"] = objs
    except Exception:
        pass

    state["topic_slug"] = args.topic_slug
    update_flags(state, topic_title=os.environ.get("TOPIC_TITLE", ""))

    _topic_title = str((state.get("flags") or {}).get("topic_title") or args.topic_slug.replace("-", " "))
    state.setdefault("messages", []).append(
        HumanMessage(content=f"주제는 '{_topic_title}'로 고정. 다른 산업으로 확장하지 말고 이 주제에 한정해 최신 자료로 RAG 업데이트.")
    )

    # 그래프 워밍업(전역 캐시 사용): 필요 시 자동 (재)빌드
    _ = _get_graph_cached()

    # 그래프 빌드 성공 후 덤프 타이머 취소(운영시 불필요한 스택 덤프 완화)
    # 환경변수 FAULT_DUMP_CANCEL_AFTER_BUILD=0 이면 유지함.
    try:
        _cancel_after_build = (_os_for_fault.getenv("FAULT_DUMP_CANCEL_AFTER_BUILD") or "1").strip().lower()
        if _cancel_after_build not in ("0", "false", "off", "no"):
            # 예약된 마지막 덤프(파일/콘솔 모두) 취소. 파일 핸들은 유지(재예약 가능).
            faulthandler.cancel_dump_traceback_later()
            logger.debug("[FAULT] dump_traceback_later cancelled after build_graph()")
    except Exception:
        pass

    logger.info(
        "Application started (config.DOC_MODE=%s, iteration_count=%s, agent_role=%s)",
        getattr(config.CFG, "DOC_MODE", None),
        getattr(config.CFG, "ITERATION_COUNT", None),
        getattr(config.CFG, "AGENT_ROLE", None),
    )
    
    logger.info("ENV CHECK: DIRECT_QA=%r ALLOW_SUMMARY=%r", os.getenv("DIRECT_QA"), os.getenv("ALLOW_SUMMARY"))

    logger.debug(
        "[CFG] AUTO_WRITE_AFTER_RAG=%s AUTO_WRITE_DURING_RESEARCH=%s DIRECT_QA=%s",
        getattr(config.CFG, "AUTO_WRITE_AFTER_RAG", None),
        getattr(config.CFG, "AUTO_WRITE_DURING_RESEARCH", None),
        getattr(config.CFG, "DIRECT_QA", None),
    )

    if os.getenv("SHOW_CONSOLE_BANNER", "1").strip().lower() in ("1","true","yes","on"):
        human_print(
            "┌─ Console: "
            f"HUMAN_LOGS={int(bool(getattr(config.CFG,'HUMAN_LOGS',False)))}, "
            f"VERBOSE={int(bool(getattr(config.CFG,'HUMAN_LOGS_VERBOSE',False))) } | "
            f"LOG_TOPK={getattr(config.CFG,'LOG_TOPK',3)} LOG_WRAP={getattr(config.CFG,'LOG_WRAP',88)} "
            f"LOG_DASHBOARD={int(bool(getattr(config.CFG,'LOG_DASHBOARD',True)))}"
        )

    if bool(getattr(args, "serve", False)):
        logger.info("Starting WEB server on http://%s:%s", args.host, args.port)
        # 콘솔 루프 대신 웹 서버로 실행
        uvicorn.run(web_app, host=args.host, port=args.port, log_level=str(args.log_level or "info").lower())
        # uvicorn가 종료되면 여기로 돌아옵니다.
        _cancel_fault_timers_and_close()
        sys.exit(0)

    try:
        while True:
            try:
                user_input = read_user_input()
            except (EOFError, KeyboardInterrupt):
                logger.info("Interrupt received. Attempting to save final state...")
                try:
                    save_state(current_path, state)
                except Exception as se:
                    logger.exception("final save_state failed: %s", se)
                finally:
                    _cancel_fault_timers_and_close()  # ← 종료 전 정리
                    logger.info("Goodbye!")
                break

            if not str(user_input).strip():
                logger.warning("빈 입력 수신. 도움말은 'help' 또는 '?'")
                continue

            if user_input.lower() in ["exit", "quit", "q"]:
                logger.info("Exit command received. Saving state...")
                try:
                    save_state(current_path, state)
                except Exception as se:
                    logger.exception("final save_state failed: %s", se)
                finally:
                    _cancel_fault_timers_and_close()  # ← 종료 전 정리
                    logger.info("Goodbye!")
                break

            _u = user_input.strip().lower()
            if _u in ("build: report", "report build", "보고서 빌드", "최종 보고서 생성"):
                slug = str(state.get("topic_slug") or "default")
                outline_fname = state.get("outline_fname") or ("outline_report.md" if config.DOC_MODE == "report" else "outline.md")
                try:
                    out_path, missing = build_final_report(
                        topic_slug=slug,
                        outline_fname=outline_fname,
                        mode=config.DOC_MODE,
                        root_dir=str(current_path)
                    )
                    update_flags(state)
                    state["last_saved_path"] = out_path
                    if missing:
                        logger.warning("Report built with missing sections: %s", ", ".join(missing))
                    print(f"\n[REPORT] 생성 완료 → {out_path}")
                    if missing:
                        print(f"[REPORT] 미수록 섹션({len(missing)}): " + ", ".join(missing))
                        # 추가 디버그: 누락 섹션 제목으로 실제 파일 존재 여부 점검
                        try:
                            _slug = slug
                            _mode = config.DOC_MODE
                            _root = str(current_path)
                            _found_tips: list[str] = []
                            for _title in missing:
                                _p = find_section_path(_title, topic_slug=_slug, mode=_mode, root_dir=_root)
                                if _p:
                                    _found_tips.append(f"  - '{_title}' → 파일 존재: {_p}")
                                else:
                                    _found_tips.append(f"  - '{[_title][0]}' → 파일 없음(작성 필요)")
                            if _found_tips:
                                print("[REPORT][HINT] 섹션 파일 탐색 결과:")
                                print("\n".join(_found_tips))
                        except Exception as _e:
                            logger.debug("find_section_path hint failed: %s", _e)
                except Exception as e:
                    logger.exception("build_final_report failed: %s", e)
                    print(f"[REPORT][ERROR] {e}")
                continue

            # ✅ v2026-03-03-E: console input must go through run_once (FastCommandHandler first, LLM 0회 보장)
            state = run_once(state, user_input, recursion_limit=args.recursion_limit)
            human_print(_last_ai_text(state) or "OK")

            # (optional) keep lightweight debug logs here (reachable)
            logger.debug("MESSAGE COUNT = %s", len(state.get("messages", [])))
            tail = [
                (getattr(t, "agent", None), getattr(t, "done", None), getattr(t, "description", None))
                for t in state.get("task_history", [])
            ][-3:]
            logger.debug("TASKS tail = %s", tail)
            try:
                flags = {
                    "agent_role": state.get("agent_role"),
                    "iteration_count": state.get("iteration_count"),
                    "research_round": state.get("research_round"),
                    "research_loop_active": state.get("research_loop_active"),
                    "has_plan": bool((state.get("research_plan") or {}).get("objective")),
                }
                logger.debug("RESEARCH flags = %s", flags)
            except Exception as e:
                logger.exception("Failed to log research flags: %s", e)
            logger.debug("last_saved_path: %s", state.get("last_saved_path"))
            # run_once 내부에서 save_state/current_path 반영까지 수행하므로 여기서 중복 저장하지 않음
            continue

    except Exception:
        # 최상위 안전망
        logger.exception("Fatal error in main loop")
        _cancel_fault_timers_and_close()  # ← 비정상 종료 시에도 정리
        raise
    finally:
        # 정상 종료 루트에서도 타이머/파일 정리 보장
        _cancel_fault_timers_and_close()
